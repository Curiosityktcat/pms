import io
import os
from docx import Document

TEMPLATE = os.path.abspath(os.path.join(
    os.path.dirname(__file__), "..", "..",
    "医院模板", "文件汇总", "4.代理协议模板（自行采购-院内竞选项目）.docx"
))


def _replace_runs(paragraph, replacements):
    """在段落内逐 run 替换占位符，保留 run 的格式与换行符(<w:br/>)。
    适用于占位符不跨 run 的情况（本模板均满足）。"""
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        for old, new in replacements:
            if old in text:
                text = text.replace(old, new)
        if text != run.text:
            run.text = text


def generate(project, agency_name, *, agency_address="", officer_name="",
             officer_phone="0832-2256120", sign_date=""):
    """按模板生成委托代理协议 Word，返回 (BytesIO, filename)。"""
    doc = Document(TEMPLATE)

    project_name = (project.name or "").strip()
    officer_name = (officer_name or project.officer or "").strip()

    # 顺序敏感：先替换更长的占位符，避免误伤
    replacements = [
        ("XXXXXXX代理公司", agency_name),                       # 落款乙方(7个X)
        ("XXXXX代理公司", agency_name),                         # 受托方(5个X)
        ("内江市第一人民医院XXXXX采购项目", project_name or "内江市第一人民医院采购项目"),
        ("黄老师", officer_name or "　"),
        ("0832-2256120", officer_phone or "0832-2256120"),
        ("签订时间：", f"签订时间：{sign_date}" if sign_date else "签订时间："),
        # 乙方地址占位符 "地址：" 与 "XXXXXXX" 分属不同 run，需单独替换裸占位符。
        # 此规则须排在所有 "XXXXXXX代理公司" 之后，避免误伤落款乙方名称。
        ("XXXXXXX", agency_address),
    ]

    for p in doc.paragraphs:
        _replace_runs(p, replacements)
    # 表格内也兜底替换一遍（本模板正文无表格，仅保险）
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_runs(p, replacements)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"委托代理协议_{project.number or project_name or '代理协议'}.docx"
    return buf, filename
