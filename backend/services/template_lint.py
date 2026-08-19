# -*- coding: utf-8 -*-
"""模板体检：上传那一刻就把问题报出来，而不是等出了稿发现一片空白。

pdt 第六节明写：「传进系统，让它解析一遍：**认不出的占位符要当场列给你看**，
而不是等生成出来一片空白才发现。」这份就是那一步。

查的每一条都是真踩过的坑，不是设想：
  · ☑ ☐ 中文字体里没有字形，用了转 PDF 时整行中文会跟着消失
  · {%tr for %} 和 {%tr endfor %} 不成对 / 写在同一行 → 循环体被整行删掉
  · 占位符被 Word 拆到多个 run 里 → 系统认不出来
  · 名字带标点空格、和字段字典撞车
  · 正文里原样打出模板语法 → 被当成代码执行，整份模板读不了
"""
import re

BAD_BOXES = {"☑": "U+2611", "☐": "U+2610", "☒": "U+2612"}
GOOD_HINT = "请改用 ■ □ √ × ● ○ 这些 GB2312 里就有的符号"
# 名字里不许出现的字符（pdt 第四条）。用字符类列举，别在字面量里塞引号。
NAME_BAD = re.compile(
    '[' + re.escape('，,。．；;：:、（）()[]【】\u201c\u201d\u2018\u2019"\'?？!！')
    + r'\s　]')


def _all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for it in cell.tables:          # 嵌套表
                    for r2 in it.rows:
                        for c2 in r2.cells:
                            for p in c2.paragraphs:
                                yield p


def lint(path, known_names=None):
    """返回 (errors, warnings, names)。errors 非空就别让它上线。"""
    from docx import Document
    errors, warns = [], []
    try:
        doc = Document(path)
    except Exception as e:                                   # noqa: BLE001
        return [f"这个文件打不开，可能不是 .docx：{e}"], [], []

    paras = list(_all_paragraphs(doc))
    full = "\n".join("".join(r.text for r in p.runs) for p in paras)

    # ① 不能用的勾选符号
    for ch, code in BAD_BOXES.items():
        if ch in full:
            errors.append(f"用了 {ch}（{code}）——中文字体里没有这个字形，"
                          f"转 PDF 时同一行的中文会跟着消失。{GOOD_HINT}")

    # ② 占位符被拆散：整段文本里有 {{ 但单个 run 里凑不出完整的
    for p in paras:
        joined = "".join(r.text for r in p.runs)
        if "{{" not in joined and "{%" not in joined:
            continue
        if joined.count("{{") != joined.count("}}"):
            errors.append(f"占位符没闭合：{joined.strip()[:50]}")
        if joined.count("{%") != joined.count("%}"):
            errors.append(f"标签没闭合：{joined.strip()[:50]}")

    # ③ 循环标记成对，且各占一行
    for kind in ("tr", "tc", "p"):
        n_for = len(re.findall(r"\{%" + kind + r" for ", full))
        n_end = len(re.findall(r"\{%" + kind + r" endfor", full))
        if n_for != n_end:
            errors.append(f"{{%{kind} for %}} 有 {n_for} 个，"
                          f"{{%{kind} endfor %}} 有 {n_end} 个，不成对")
    for p in paras:
        joined = "".join(r.text for r in p.runs)
        if re.search(r"\{%\w* for ", joined) and re.search(r"\{%\w* endfor", joined):
            errors.append(f"for 和 endfor 写在同一段了，会被整段删掉："
                          f"{joined.strip()[:50]}")

    # ④ 占位符名字
    names = set()
    # {{ 名字 }} 和 {{p 名字 }} 都要认。指令字母（p/r/tr/tc）跟在 {{ 后面、
    # 与名字之间有个空格——`[a-z]?` 那种写法会把「包」当成名字的一部分丢掉，
    # 反过来把指令字母当成名字（第一次跑出来报「{{ .xxx }}」这种鬼东西）。
    for m in re.finditer(r"\{\{\s*(?:(?:tr|tc|p|r)\s+)?([^\s|}]+)", full):
        n = m.group(1).strip()
        if n:
            names.add(n)
    loop_vars = {n.split(".", 1)[0] for n in names if "." in n}
    plain = {n for n in names if "." not in n}
    for n in plain:
        if NAME_BAD.search(n):
            errors.append(f"占位符名字里有标点或空格：「{n}」——名字里只能用字和数字")
        if len(n) > 24:
            warns.append(f"占位符名字太长，不好认：「{n}」")
    # 循环变量必须有对应的 for
    for v in loop_vars:
        if not re.search(r"\{%\w* for " + re.escape(v) + r" in ", full):
            errors.append(f"用了 {{{{ {v}.xxx }}}}，但没有对应的 "
                          f"{{%tr for {v} in ... %}}")

    # ⑤ 和字段字典对不上
    if known_names is not None:
        unknown = plain - set(known_names)
        for n in sorted(unknown):
            warns.append(f"「{n}」不在字段字典里——出稿时会留空，"
                          f"确认是要人手填的自由段落再忽略这条")

    # ⑥ 正文里原样打出模板语法（pdt 第七条）
    if re.search(r"[「『\u201c][^\u201d」』]{0,10}\{[{%]", full):
        warns.append("正文里像是原样写了模板语法举例——它会被当成代码执行，"
                     "要提语法请用文字描述或把大括号换成全角")

    return errors, warns, sorted(plain)
