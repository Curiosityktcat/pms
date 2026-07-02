import io, os, json
from docx import Document

TEMPLATE = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "..", "..", "医院模板", "1.盖章文件", "4.（打印1份盖章）结果确认函.docx"
)

def _cn_amount(amount: float) -> str:
    """将数字金额转为中文大写（简单版）"""
    digits = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']
    if amount == 0:
        return '零元整'
    result = f'{amount:.2f}'
    integer_part, decimal_part = result.split('.')
    int_num = int(integer_part)
    # 简单处理：直接用格式化
    try:
        yi = int_num // 100000000
        wan = (int_num % 100000000) // 10000
        qian = int_num % 10000
        parts = []
        if yi: parts.append(f'{yi}亿')
        if wan: parts.append(f'{wan}万')
        if qian: parts.append(f'{qian}')
        cn = ''.join(parts) + '元'
        fen = int(decimal_part)
        if fen == 0:
            cn += '整'
        elif decimal_part[0] != '0':
            cn += f'{digits[int(decimal_part[0])]}角{digits[int(decimal_part[1])] if decimal_part[1]!="0" else ""}分'
        else:
            cn += f'零{digits[int(decimal_part[1])]}分'
        return cn
    except Exception:
        return f'{amount}元'


def _set_cell(table, row, col, text):
    try:
        cell = table.cell(row, col)
        cell.text = text
    except Exception:
        pass


def _append_right(cell, text):
    """在单元格末尾追加一个右对齐段落（采购人盖章/日期放末行靠右），字体沿用正文首段。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    base = cell.paragraphs[0] if cell.paragraphs else None
    para = cell.add_paragraph()
    if base is not None:
        para.style = base.style
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(text)
    if base is not None and base.runs:
        s = base.runs[0]
        if s.font.size:
            run.font.size = s.font.size
        if s.font.name:
            run.font.name = s.font.name
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), s.font.name)
    return para


def generate(result, project):
    doc = Document(TEMPLATE)
    t = doc.tables[0]

    pname = project.name if project else ""
    pnumber = project.number if project else ""

    # 第二次及以后的竞选，项目名称后追加「（第二次）」等（与采购公告格式一致）
    ROUND_CN = ["", "一", "二", "三", "四", "五"]
    round_label = ""
    if result.round_number and result.round_number > 1:
        cn = ROUND_CN[min(result.round_number, len(ROUND_CN) - 1)]
        round_label = f"（第{cn}次）"

    try:
        packages = json.loads(result.packages_json or "[]")
    except Exception:
        packages = []

    # Row0: 项目名称
    _set_cell(t, 0, 1, f"内江市第一人民医院{pname}采购项目{round_label}")
    # Row1: 项目编号 | 采购方式
    _set_cell(t, 1, 1, pnumber)
    _set_cell(t, 1, 3, result.procurement_method or "院内竞选")
    # Row2: 招标代理机构
    _set_cell(t, 2, 3, result.agency_name or "")
    # Row3: 竞选时间
    _set_cell(t, 3, 3, result.bid_time or "")
    # Row4: 备注
    _set_cell(t, 4, 1, result.notes or "此结果为评审委员会评审结果")

    # Row5: 采购人意见（动态拼接包结果）
    lines = ["经评审委员会的综合评审：", "", ""]
    for i, pkg in enumerate(packages, start=1):
        result_val = pkg.get("result", "废标")
        winner = pkg.get("winner", "")
        amount = pkg.get("amount", 0)
        try:
            amount_f = float(amount)
        except Exception:
            amount_f = 0
        amount_cn = pkg.get("amount_cn", "") or _cn_amount(amount_f)

        chk_deal = "☑" if result_val == "成交" else "□"
        chk_fail = "☑" if result_val == "废标" else "□"
        lines.append(f"采购包{i}：{chk_deal}成交{chk_fail}废标")
        if result_val == "成交":
            lines.append(f"中标人：{winner}。")
            if pkg.get("unit_price_attached"):
                lines.append("中标单价：详见附件。")
            else:
                lines.append(f"中标金额：￥{amount_f:.2f}（大写：{amount_cn}）。若采购内容过多，请附报价详单")
        else:
            lines.append(f"废标原因：{pkg.get('note', '')}")
        lines.append("")

    lines.append("")
    lines.append("")

    _set_cell(t, 5, 1, "\n".join(lines))
    # 采购人盖章 + 确认日期 —— 末行靠右（各自独立右对齐段落）
    confirm_date = result.confirm_date or ""
    cell5 = t.cell(5, 1)
    _append_right(cell5, "内江市第一人民医院（盖章）")
    if confirm_date:
        _append_right(cell5, confirm_date)

    buf = io.BytesIO()
    from services.docx_utils import strip_highlight
    strip_highlight(doc)                 # 清除模板黄色高亮占位印记
    doc.save(buf)
    buf.seek(0)
    round_suffix = f"第{result.round_number}次" if result.round_number and result.round_number > 1 else ""
    filename = f"采购结果确认函_{pnumber}{round_suffix}.docx"
    return buf, filename
