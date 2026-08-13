"""rd-web 合同审签「AI 自动填写」共享逻辑。

工具页（rdweb_contract_api）与合同管理推送（contract_api）共用：
附件文本提取 → LLM 逐字抽取 → 甲方兜底 / 代理机构乙方补全 / 合同名称拼项目名。
"""
import os

# 甲方恒为本院，识别不到时兜底填写
HOSPITAL_PARTY_A = {
    "合同甲方":       "内江市第一人民医院",
    "甲方法定代表人": "谢晓阳",
    "甲方联系电话":   "0832-2256120",
    "甲方地址":       "四川省内江市市中区沱中路41号、汉安大道西段1866号",
}

# 代理机构兜底信息（agencies 表缺数据时使用；与基础数据维护页同源）
FALLBACK_AGENCIES = [
    ("四川知行招标代理有限公司",       "银钒霖", "0832-2029668", "四川省内江市东兴区万达中心2112-2115号"),
    ("内江中洲工程项目管理有限公司",   "周丽萍", "13882007116",  "内江市东兴区中兴路1104号5幢1单元207号（上海花园）"),
    ("四川中锦招标代理有限公司",       "何根炜", "0832-2242423", "内江市汉安大道西段927号40幢2楼1号、2号"),
    ("四川尚璟招标代理有限责任公司",   "余尚英", "0832-2267533", "内江市东兴区兰桂大道222号负一层3号、4号、5号、6号、7号"),
    ("四川华询工程管理有限责任公司",   "唐国英", "0832-2035990", "内江市东兴区北环路四季康城12号楼【幢】无单元2层2-5号"),
    ("四川三盈招标代理有限公司",       "余勇",   "0832-2111314", "内江市东兴区胜利路666号1栋2单元A区7层6号（汉安大道传化广场锦城A区）"),
    ("内江市川交公路勘察设计有限公司", "田玉冲", "17313721737 0832-2223383", "四川省内江市东兴区梧桐路1号二幢2单元27楼23号"),
    ("内江市政府采购中心",             "卿伟",   "0832-2048628", "四川省内江市东兴区兰桂大道377号"),
]

AUTOFILL_SYSTEM = """你是采购合同信息抽取助手。用户会粘贴一份合同的首页（或前几页）文字，
请从中抽取合同审签单所需字段，严格只输出一个 JSON 对象，键为：
合同名称、合同编码、项目名称及包号、归口管理科室、合同金额、
合同甲方、甲方法定代表人、甲方联系电话、甲方地址、
合同乙方、乙方法定代表人、乙方联系电话、乙方地址。

规则：
1. 值必须逐字取自原文，不要改写、不要补全；原文没有的字段填空字符串 ""。
2. 合同金额保留原文写法（含币种符号/大写均可），若同时有大小写以数字小写为准。
3. 甲方通常是采购人/买方/需方（医院），乙方是供应商/卖方/供方。
4. 项目名称及包号：合同中提到的采购项目名称，含包号则一并带上。
5. 若这是一份委托代理协议：注意看代理服务费由谁支付——约定由供应商/中标（成交）
   人支付、医院（甲方/采购人）不支付费用的，合同金额填斜杠 "/"；
   约定由医院支付的才填实际金额。
6. 只输出 JSON，不要任何解释文字。"""

FIELD_KEYS = [
    "合同名称", "合同编码", "项目名称及包号", "归口管理科室", "合同金额",
    "合同甲方", "甲方法定代表人", "甲方联系电话", "甲方地址",
    "合同乙方", "乙方法定代表人", "乙方联系电话", "乙方地址",
    "合同类别", "经办人",
]

# ── 附件文本提取 ─────────────────────────────────────────────────
EXTRACT_MAX_CHARS = 8000   # 审签字段都在前几页，够用即止


def _docx_text(path):
    from docx import Document
    d = Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for tb in d.tables:
        for r in tb.rows:
            parts.append("　".join(c.text.strip() for c in r.cells if c.text.strip()))
    return "\n".join(parts)


def _pdf_text(path):
    import fitz
    text = []
    with fitz.open(path) as pdf:
        for page in pdf.pages(0, min(6, pdf.page_count)):
            text.append(page.get_text())
            if sum(len(t) for t in text) > EXTRACT_MAX_CHARS:
                break
    return "\n".join(text)


def _ocr_text(path):
    """扫描版 PDF / 图片：走本机 OCR 服务（免费传统引擎）。"""
    import requests as rq
    ocr_base = os.environ.get("PMS_OCR_URL", "http://127.0.0.1:8118").rstrip("/")
    with open(path, "rb") as f:
        r = rq.post(f"{ocr_base}/ocr_classic",
                    files={"file": (os.path.basename(path), f)}, timeout=300)
    r.raise_for_status()
    return r.json().get("markdown") or ""


def extract_file_text(path):
    """尽力从合同附件里取出文字；失败抛 RuntimeError（带用户可读原因）。"""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".docx":
            text = _docx_text(path)
        elif ext == ".doc":
            from services.procurement_doc_gen import _to_docx
            text = _docx_text(_to_docx(path))
        elif ext == ".pdf":
            text = _pdf_text(path)
            if len(text.strip()) < 60:      # 无文字层 → 扫描件，转 OCR
                text = _ocr_text(path)
        elif ext in (".jpg", ".jpeg", ".png"):
            text = _ocr_text(path)
        else:
            raise RuntimeError(f"{ext} 附件不支持内容识别，请粘贴合同文字")
    except RuntimeError:
        raise
    except Exception as e:
        raise RuntimeError(f"附件内容读取失败：{str(e)[:150]}")
    text = (text or "").strip()
    if len(text) < 20:
        raise RuntimeError("附件里没读到有效文字（可能是低质量扫描件），请粘贴合同文字")
    return text[:EXTRACT_MAX_CHARS]


def complete_party_b_from_agency(out):
    """乙方是已知代理机构时，用基础数据补全缺失的法人/电话/地址。"""
    name = (out.get("合同乙方") or "").strip()
    if not name:
        return
    info = None
    try:
        from models.agency import Agency
        for ag in Agency.query.all():
            if ag.name and (ag.name in name or name in ag.name):
                info = (ag.name, ag.legal_rep or "", ag.phone or "", ag.address or "")
                break
    except Exception:
        pass
    if info is None or not all(info[1:]):
        for row in FALLBACK_AGENCIES:
            if row[0] in name or name in row[0]:
                # 库里查到但字段有缺的，用兜底表逐项补
                info = tuple(info[i] if info and info[i] else row[i] for i in range(4)) if info else row
                break
    if not info:
        return
    out["合同乙方"] = info[0]
    for key, val in (("乙方法定代表人", info[1]),
                     ("乙方联系电话",   info[2]),
                     ("乙方地址",       info[3])):
        if not out.get(key) and val:
            out[key] = val


def autofill_fields(text, *, usage_ctx=None, operator=""):
    """合同文字 → 审签字段 dict。须在 app 上下文内调用；失败抛 RuntimeError。"""
    from services.llm_client import chat_json
    try:
        data = chat_json(AUTOFILL_SYSTEM, text[:EXTRACT_MAX_CHARS],
                         temperature=0.1, max_tokens=1200, usage_ctx=usage_ctx)
    except Exception as e:
        raise RuntimeError(f"模型识别失败：{str(e)[:200]}")
    if not isinstance(data, dict):
        raise RuntimeError("模型未返回有效 JSON")

    out = {k: str(data.get(k) or "").strip() for k in FIELD_KEYS if k in data}

    # 合同名称 = 合同名称 + 项目名称（项目名尚未包含在合同名里时追加）
    cname, pname = out.get("合同名称", ""), out.get("项目名称及包号", "")
    if cname and pname and pname not in cname:
        out["合同名称"] = f"{cname}（{pname}）"

    # 甲方恒为本院：识别缺失的一律用医院默认信息补齐
    for k, v in HOSPITAL_PARTY_A.items():
        if not out.get(k):
            out[k] = v

    # 乙方为已知代理机构（如代理协议）时，从基础数据补全法人/电话/地址
    complete_party_b_from_agency(out)

    out.setdefault("合同类别", "采购部合同")
    if not out.get("经办人"):
        out["经办人"] = operator or "黄新博"
    return out
