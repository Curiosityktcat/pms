"""更正公告 Word 生成（6.3）。

无医院模板，按内江一院官网真实更正公告体例（三段式）用 python-docx 从零构建：
  标题：{项目名}更正公告（第N次）
  一、项目基本情况（原公告编号/名称/发布日期）
  二、更正信息（更正事项/原因/内容[逐条更正前→更正后 或 详见附件]/其他内容不变/更正日期）
  三、联系方式（采购人固定 + 代理机构）
"""
import io
import json
import datetime

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROUND_CN = ["", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]

# 采购人固定信息（与采购公告模板 1.9 联系方式一致）
PURCHASER_LINES = [
    "采购人：内江市第一人民医院",
    "地址：四川省内江市市中区沱中路41号、汉安大道西段1866号",
    "邮编：641000",
    "联系人：黄老师",
    "联系电话：0832-2256120",
]


def _cn(n):
    return ROUND_CN[n] if 0 < n < len(ROUND_CN) else str(n)


def _para(doc, text, size=14, bold=False, center=False, indent=True, font="仿宋_GB2312"):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    return p


def _today_cn():
    d = datetime.date.today()
    return f"{d.year}年{d.month:02d}月{d.day:02d}日"


def generate(project, ann, agency_name):
    """生成更正公告 Word，返回 BytesIO。"""
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(3.17)
        s.right_margin = Cm(3.17)

    pname = project.name or ""
    pnumber = project.number or ""
    round_no = ann.round_number or 1
    seq = ann.corr_seq or 1
    # 项目名带次数后缀（第二轮及以后），与采购公告标题习惯一致
    name_with_round = pname if round_no <= 1 else f"{pname}（第{_cn(round_no)}次）"

    # ── 标题 ──
    _para(doc, f"{name_with_round}更正公告（第{_cn(seq)}次）",
          size=16, bold=True, center=True, font="黑体")
    doc.add_paragraph()

    # ── 一、项目基本情况 ──
    _para(doc, "一、项目基本情况", size=14, bold=True, indent=False, font="黑体")
    _para(doc, f"1.原公告的采购项目编号：{pnumber}")
    _para(doc, f"2.原公告的采购项目名称：{name_with_round}")

    # ── 二、更正信息 ──
    _para(doc, "二、更正信息", size=14, bold=True, indent=False, font="黑体")
    _para(doc, f"更正事项：{ann.corr_scope or '采购公告'}")
    if (ann.corr_reason or "").strip():
        _para(doc, f"更正原因：{ann.corr_reason.strip()}")
    _para(doc, "更正内容：")
    if ann.corr_in_attachment:
        _para(doc, "更正内容较多，详见本公告附件。")
    else:
        try:
            items = json.loads(ann.corr_items_json or "[]")
        except Exception:
            items = []
        for i, it in enumerate(items, 1):
            label = (it.get("item") or "").strip()
            before = (it.get("before") or "").strip()
            after = (it.get("after") or "").strip()
            prefix = f"{i}." if len(items) > 1 else ""
            if label:
                _para(doc, f"{prefix}{label}")
                _para(doc, f"更正前：{before}")
            else:
                _para(doc, f"{prefix}更正前：{before}")
            _para(doc, f"更正后：{after}")
    _para(doc, "其他内容不变。")
    corr_date = _fmt_date(ann.confirmed_at) if ann.confirmed_at else _today_cn()
    _para(doc, f"更正日期：{corr_date}")

    # ── 三、联系方式 ──
    _para(doc, "三、凡对本次公告内容提出询问，请按以下方式联系。",
          size=14, bold=True, indent=False, font="黑体")
    for line in PURCHASER_LINES:
        _para(doc, line, indent=False)
    _para(doc, f"代理机构：{agency_name}", indent=False)
    if ann.agency_address:
        _para(doc, f"地址：{ann.agency_address}", indent=False)
    if ann.agency_contact:
        _para(doc, f"联系人：{ann.agency_contact}", indent=False)
    if ann.agency_contact_phone:
        _para(doc, f"联系电话：{ann.agency_contact_phone}", indent=False)

    # ── 落款 ──
    doc.add_paragraph()
    p = _para(doc, agency_name, indent=False)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2 = _para(doc, _fmt_date(ann.confirmed_at) if ann.confirmed_at else _today_cn(), indent=False)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _fmt_date(iso):
    """'2026-06-10T09:00:00' → '2026年06月10日'；解析失败返回原串。"""
    try:
        d = datetime.date.fromisoformat((iso or "")[:10])
        return f"{d.year}年{d.month:02d}月{d.day:02d}日"
    except Exception:
        return iso or ""


def get_filename(project, ann):
    seq = ann.corr_seq or 1
    return f"更正公告_{project.number}（第{_cn(seq)}次）.docx"
