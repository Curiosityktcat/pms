"""归档「一键打印资料」：按轮次顺序把项目各要件 Word 合并成一份可打印文档。

每一轮（第1次、第2次……）按以下顺序汇总：
  1. 采购文件确认函（《院内竞选文件内容确认表》）
  2. 采购文件封面（打印两份）
  3. 授权函
  4. 采购结果确认函（该轮全部废标则省略；存在成交包即收录）

各子文档之间插入分页符，整体返回单个 .docx（前端预览后直接打印）。
"""
import json
import os
from io import BytesIO

from docx import Document
from docxcompose.composer import Composer

from models import db
from models.procurement_round import ProcurementRound
from models.procurement_doc_attachment import ProcurementDocAttachment
from models.auth_letter_record import AuthLetterRecord
from models.procurement_result import ProcurementResult
from models.people import People
from services.project import get_agency_name


def _round_numbers(pid):
    """汇总该项目出现过的全部轮次号（跨轮次表/授权函/采购结果/采购文件附件）。"""
    nums = set()
    for model in (ProcurementRound, AuthLetterRecord, ProcurementResult):
        rows = db.session.execute(
            db.select(model.round_number).filter_by(project_id=pid)
        ).scalars().all()
        nums.update(int(n or 1) for n in rows)
    rows = db.session.execute(
        db.select(ProcurementDocAttachment.round_number)
        .filter_by(project_id=pid, kind="doc")
    ).scalars().all()
    nums.update(int(n or 1) for n in rows)
    return sorted(nums) or [1]


def _find_person(name):
    name = (name or "").strip()
    if not name:
        return None
    return db.session.execute(db.select(People).filter_by(name=name)).scalars().first()


def _confirm_date_cn(rnd):
    """该轮经办人确认采购文件的当天 → 中文「YYYY年M月D日」。"""
    raw = getattr(rnd, "doc_confirmed_at", "") if rnd else ""
    if not raw:
        return ""
    try:
        import datetime
        dt = datetime.datetime.fromisoformat(raw)
        return f"{dt.year}年{dt.month}月{dt.day}日"
    except Exception:
        return ""


def _content_confirm_doc(project, agency_name, rno):
    """采购文件确认函（内容确认表）：仅在该轮采购文件已确认时收录。"""
    rnd = db.session.execute(
        db.select(ProcurementRound).filter_by(project_id=project.id, round_number=rno)
    ).scalars().first()
    if not (rnd and rnd.doc_confirmed):
        return None
    docs = db.session.execute(
        db.select(ProcurementDocAttachment)
        .filter_by(project_id=project.id, kind="doc", round_number=rno)
        .order_by(ProcurementDocAttachment.id)
    ).scalars().all()
    file_hashes = [(d.original_name, d.sha256) for d in docs if d.sha256]
    contact = rnd.doc_agency_contact or project.doc_agency_contact or ""
    phone = rnd.doc_agency_phone or project.doc_agency_phone or ""
    from services.content_confirm_word import generate
    buf, _ = generate(project, agency_name, file_hashes=file_hashes,
                      contact_person=contact, contact_phone=phone,
                      confirm_date=_confirm_date_cn(rnd))
    return Document(buf)


def _bid_cover_doc(project, agency_name, rno):
    """采购文件封面（每次调用生成独立副本，便于打印两份）。"""
    rnd = db.session.execute(
        db.select(ProcurementRound).filter_by(project_id=project.id, round_number=rno)
    ).scalars().first()
    from services.bid_cover_word import generate
    buf, _ = generate(project, agency_name,
                      compile_date=_confirm_date_cn(rnd), round_number=rno)
    return Document(buf)


def _auth_letter_doc(project, agency_name, rno):
    """授权函：按该轮已保存的授权函记录重新生成（人员从人员库按姓名反查）。"""
    rec = db.session.execute(
        db.select(AuthLetterRecord)
        .filter_by(project_id=project.id, round_number=rno)
        .order_by(AuthLetterRecord.id.desc())
    ).scalars().first()
    if not rec:
        return None
    supervisor = _find_person(rec.supervisor_name)
    rep_names = [n for n in (rec.representative_names or "").replace(",", "、").split("、") if n.strip()]
    representatives = [p for p in (_find_person(n) for n in rep_names) if p]
    if not supervisor or not representatives:
        return None
    from services import auth_letter as svc
    tmp_path = svc.generate(project, supervisor, representatives, agency_name,
                            round_number=rno, bid_time_override=rec.bid_time or "")
    try:
        doc = Document(tmp_path)
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass
    return doc


def _result_doc(project, rno):
    """采购结果确认函：该轮全部废标则不收录，存在成交包即收录。"""
    res = db.session.execute(
        db.select(ProcurementResult)
        .filter_by(project_id=project.id, round_number=rno)
        .order_by(ProcurementResult.id.desc())
    ).scalars().first()
    if not res:
        return None
    try:
        items = json.loads(res.packages_json or "[]")
    except Exception:
        items = []
    if not any(it.get("result") == "成交" for it in items):
        return None
    from services.procurement_result_word import generate
    buf, _ = generate(res, project)
    return Document(buf)


def build_print_bundle(project):
    """合并项目全部归档要件为单个 docx，返回 (BytesIO, manifest) 或 (None, [])。

    manifest: 收录清单 [{round, items:[名称...]}]，供前端提示。
    """
    agency_name = get_agency_name(project.agency_code) or project.agency_code or ""
    docs = []
    manifest = []
    for rno in _round_numbers(project.id):
        items = []
        cc = _content_confirm_doc(project, agency_name, rno)
        if cc is not None:
            docs.append(cc)
            items.append("采购文件确认函")
        # 采购文件封面：打印两份
        docs.append(_bid_cover_doc(project, agency_name, rno))
        docs.append(_bid_cover_doc(project, agency_name, rno))
        items.append("采购文件封面 ×2")
        al = _auth_letter_doc(project, agency_name, rno)
        if al is not None:
            docs.append(al)
            items.append("授权函")
        rd = _result_doc(project, rno)
        if rd is not None:
            docs.append(rd)
            items.append("采购结果确认函")
        manifest.append({"round": rno, "items": items})

    if not docs:
        return None, []

    master = docs[0]
    composer = Composer(master)
    for d in docs[1:]:
        master.add_page_break()
        composer.append(d)
    out = BytesIO()
    composer.save(out)
    out.seek(0)
    return out, manifest


# ── 归档「文件夹视图」：按轮次逐要件单独浏览/下载 ──────────────────────────

# 各要件类型（kind → 中文名），顺序与打印顺序一致
ITEM_KINDS = [
    ("content_confirm", "采购文件确认函"),
    ("bid_cover",       "采购文件封面"),
    ("auth_letter",     "授权函"),
    ("result",          "采购结果确认函"),
]

_KIND_LABEL = dict(ITEM_KINDS)


def _cn_round(n):
    return f"第{'一二三四五六七八九十'[n - 1]}次" if 1 <= n <= 10 else f"第{n}次"


def _item_available(project, rno, kind):
    """轻量判断某轮某要件是否存在（不触发文档生成，避免列目录时开销过大）。"""
    if kind == "bid_cover":
        return True   # 封面始终可生成
    if kind == "content_confirm":
        rnd = db.session.execute(
            db.select(ProcurementRound).filter_by(project_id=project.id, round_number=rno)
        ).scalars().first()
        return bool(rnd and rnd.doc_confirmed)
    if kind == "auth_letter":
        rec = db.session.execute(
            db.select(AuthLetterRecord)
            .filter_by(project_id=project.id, round_number=rno)
            .order_by(AuthLetterRecord.id.desc())
        ).scalars().first()
        if not rec:
            return False
        sup = _find_person(rec.supervisor_name)
        rep_names = [n for n in (rec.representative_names or "").replace(",", "、").split("、") if n.strip()]
        reps = [p for p in (_find_person(n) for n in rep_names) if p]
        return bool(sup and reps)
    if kind == "result":
        res = db.session.execute(
            db.select(ProcurementResult)
            .filter_by(project_id=project.id, round_number=rno)
            .order_by(ProcurementResult.id.desc())
        ).scalars().first()
        if not res:
            return False
        try:
            items = json.loads(res.packages_json or "[]")
        except Exception:
            items = []
        return any(it.get("result") == "成交" for it in items)
    return False


def list_archive_tree(project):
    """归档「文件夹视图」：统一为 [{folder, items:[{name,url,preview_url,size}]}]。

    含 ① 按轮次即时生成的要件（采购文件确认函/封面/授权函/采购结果确认函），
    ② 项目在各业务模块下上传/生成的真实文件资料（采购需求/文件/评审/结果附件/
       合同/公告/询价等，见 services.archive_files）。仅返回有内容的文件夹。
    """
    pid = project.id
    folders = []

    # ① 按轮次生成的要件
    for rno in _round_numbers(pid):
        items = []
        for kind, label in ITEM_KINDS:
            if _item_available(project, rno, kind):
                items.append({
                    "name": f"{label}.docx",
                    "size": 0,
                    "url": f"/api/archive/{pid}/item?round={rno}&kind={kind}&download=1",
                    "preview_url": f"/api/archive/{pid}/item?round={rno}&kind={kind}",
                })
        if items:
            folders.append({"folder": f"{_cn_round(rno)} · 生成要件", "items": items})

    # ② 各模块上传/生成的真实文件资料
    from services.archive_files import collect_material_folders
    folders.extend(collect_material_folders(project))
    return folders


def build_item(project, rno, kind):
    """生成单个归档要件，返回 (BytesIO, 中文名) 或 (None, None)。"""
    agency_name = get_agency_name(project.agency_code) or project.agency_code or ""
    if kind == "content_confirm":
        doc = _content_confirm_doc(project, agency_name, rno)
    elif kind == "bid_cover":
        doc = _bid_cover_doc(project, agency_name, rno)
    elif kind == "auth_letter":
        doc = _auth_letter_doc(project, agency_name, rno)
    elif kind == "result":
        doc = _result_doc(project, rno)
    else:
        return None, None
    if doc is None:
        return None, None
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out, _KIND_LABEL.get(kind, kind)
