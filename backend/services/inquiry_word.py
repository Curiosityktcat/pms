"""
生成询/议价邀请函 Word 文档（python-docx）

体例参照样板《内江市第一人民医院询议价邀请函》：
    标题（询/议价邀请函）
    一、项目名称
    二、项目编号        ← 取自关联项目，自动
    三、项目细则和限价   ← 默认取项目立项内容，可改
    四、相关要求        ← 经办人填写
    五、资料递交        ← 截止日期 + 收件邮箱
    六、其他信息（地址 / 联系人=经办人 / 电话）
    落款：内江市第一人民医院 + 当日日期
"""
from io import BytesIO
import datetime

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 采购部固定信息
RECEIVE_EMAIL = "njyycgbxjh@163.com"
DEPT_ADDRESS  = "内江市汉安大道西段1866号，内江市第一人民医院新区全科医师楼三楼采购部。"
DEPT_PHONE    = "0832-2256120"


def generate_inquiry_word(letter, project_name: str, project_number: str,
                          officer: str = "") -> BytesIO:
    """
    letter: InquiryLetter 对象（含 type / detail / requirements / deadline 等）
    project_name / project_number: 取自关联项目
    officer: 经办人（取自项目）
    返回包含 Word 文档内容的 BytesIO 对象
    """
    doc = Document()

    # ── 页面边距 ──────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # 正文默认字体大小
    base_size = Pt(12)

    def add_line(text="", *, bold=False, indent=True, align=None, size=base_size):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        if indent:
            p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进 2 字符
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        run.font.size = size
        run.font.bold = bold
        return p

    letter_type = getattr(letter, "type", "询价")

    # ── 标题 ──────────────────────────────────────────────────────
    title = add_line(f"内江市第一人民医院{letter_type}邀请函",
                     bold=True, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                     size=Pt(16))
    title.paragraph_format.space_after = Pt(12)

    # ── 一、项目名称 ──────────────────────────────────────────────
    add_line(f"一、项目名称：{project_name or '—'}")

    # ── 二、项目编号（自动取项目编号）────────────────────────────
    add_line(f"二、项目编号：{project_number or '—'}")

    # ── 三、项目细则和限价 ────────────────────────────────────────
    detail = (getattr(letter, "detail", "") or "").strip()
    add_line(f"三、项目细则和限价：{detail or '—'}")

    # ── 四、相关要求 ──────────────────────────────────────────────
    requirements = (getattr(letter, "requirements", "") or "").strip()
    add_line(f"四、相关要求：{requirements or '—'}")

    # ── 五、资料递交（截止日期 + 收件邮箱）────────────────────────
    deadline = (getattr(letter, "deadline", "") or "").strip()
    deadline_text = f"{deadline} 前" if deadline else "询价公告之日起规定时间内"
    add_line(
        f"五、资料递交：以上资料扫描件加盖公章（格式PDF）需于 {deadline_text} "
        f"发送至 {RECEIVE_EMAIL} 电子邮箱有效，超过截止日期则相关资料无效。"
    )

    # ── 六、其他信息 ──────────────────────────────────────────────
    add_line("六、其他信息：")
    add_line(f"地址：{DEPT_ADDRESS}")
    add_line(f"联系人：{officer or '经办人'}   电话：{DEPT_PHONE}（来电请告知采购项目）")

    # ── 落款 ──────────────────────────────────────────────────────
    doc.add_paragraph()
    sign = add_line("内江市第一人民医院", indent=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    sign.paragraph_format.line_spacing = 1.5
    today = datetime.date.today()
    add_line(f"{today.year}年{today.month}月{today.day}日",
             indent=False, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ── 保存到内存 ────────────────────────────────────────────────
    buf = BytesIO()
    from services.docx_utils import strip_highlight
    strip_highlight(doc)                 # 清除模板黄色高亮占位印记
    doc.save(buf)
    buf.seek(0)
    return buf
