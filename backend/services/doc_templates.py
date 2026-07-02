# -*- coding: utf-8 -*-
"""文档式表单的模板配置（章节/字段）+ Word 导出 + 从项目预填。

学四川政采：模板=若干章节，章节=若干字段。字段 type(text/textarea/number/date/select)、
layout(inline 行内「标签：值」/ block 长文本段落)。改字段只改这里，不动数据结构。
"""
from io import BytesIO


# ── 模板定义 ─────────────────────────────────────────────────────
TEMPLATES = {
    "procurement_demand": {
        "key": "procurement_demand",
        "name": "政府采购项目采购需求",
        "subtitle": "（采购人编制）",
        "sections": [
            {"key": "basic", "title": "一、采购项目基本情况", "fields": [
                {"key": "project_name", "label": "项目名称", "type": "text", "layout": "inline", "required": True},
                {"key": "purchaser", "label": "采购单位", "type": "text", "layout": "inline"},
                {"key": "project_no", "label": "项目编号", "type": "text", "layout": "inline"},
                {"key": "year", "label": "采购年度", "type": "text", "layout": "inline"},
                {"key": "budget", "label": "预算金额(元)", "type": "number", "layout": "inline"},
                {"key": "category", "label": "采购品目类别", "type": "select", "layout": "inline",
                 "options": ["货物", "服务", "工程"]},
                {"key": "procure_method", "label": "采购方式", "type": "select", "layout": "inline",
                 "options": ["公开招标", "竞争性谈判", "竞争性磋商", "询价", "单一来源", "框架协议", "院内竞选"]},
                {"key": "demand_dept", "label": "需求编制部门", "type": "text", "layout": "inline"},
                {"key": "compile_date", "label": "编制日期", "type": "date", "layout": "inline"},
            ]},
            {"key": "overview", "title": "二、采购需求概述", "fields": [
                {"key": "project_overview", "label": "项目概述", "type": "textarea", "layout": "block"},
                {"key": "has_related_supplier", "label": "是否涉及关联供应商", "type": "select",
                 "layout": "inline", "options": ["否", "是"]},
            ]},
            {"key": "survey", "title": "三、市场调研情况", "fields": [
                {"key": "survey_industry", "label": "相关产业发展情况", "type": "textarea", "layout": "block"},
                {"key": "survey_market", "label": "市场供给情况", "type": "textarea", "layout": "block"},
                {"key": "survey_history", "label": "同类项目历史成交情况", "type": "textarea", "layout": "block"},
                {"key": "survey_followup", "label": "后续采购安排", "type": "textarea", "layout": "block"},
                {"key": "survey_other", "label": "其他相关情况", "type": "textarea", "layout": "block"},
            ]},
            {"key": "target", "title": "四、采购标的需求", "fields": [
                {"key": "target_name", "label": "采购标的名称", "type": "text", "layout": "inline"},
                {"key": "target_qty", "label": "数量", "type": "text", "layout": "inline"},
                {"key": "tech_require", "label": "主要技术要求", "type": "textarea", "layout": "block"},
                {"key": "quality_require", "label": "质量及验收标准", "type": "textarea", "layout": "block"},
            ]},
            {"key": "business", "title": "五、商务要求", "fields": [
                {"key": "delivery_time", "label": "交付/服务时间", "type": "text", "layout": "inline"},
                {"key": "delivery_place", "label": "交付/服务地点", "type": "text", "layout": "inline"},
                {"key": "payment", "label": "付款方式", "type": "textarea", "layout": "block"},
                {"key": "warranty", "label": "售后服务要求", "type": "textarea", "layout": "block"},
            ]},
            {"key": "risk", "title": "六、风险防控措施", "fields": [
                {"key": "risk_control", "label": "风险防控措施", "type": "textarea", "layout": "block"},
            ]},
        ],
    },
    "procurement_doc": {
        "key": "procurement_doc",
        "name": "采购文件",
        "subtitle": "（采购人/代理机构编制）",
        "sections": [
            {"key": "ann", "title": "第一章 采购公告", "fields": [
                {"key": "project_intro", "label": "项目概况", "type": "textarea", "layout": "block"},
                {"key": "demand_summary", "label": "采购需求概要", "type": "textarea", "layout": "block"},
                {"key": "max_price", "label": "最高限价", "type": "text", "layout": "inline"},
                {"key": "get_doc_way", "label": "获取采购文件的方式", "type": "textarea", "layout": "block"},
                {"key": "response_deadline", "label": "响应文件递交截止时间", "type": "text", "layout": "inline"},
                {"key": "open_time", "label": "开标时间及地点", "type": "text", "layout": "inline"},
            ]},
            {"key": "qualify", "title": "第二章 供应商资格要求", "fields": [
                {"key": "qualification", "label": "供应商资格条件", "type": "textarea", "layout": "block"},
            ]},
            {"key": "require", "title": "第三章 采购需求", "fields": [
                {"key": "tech_spec", "label": "技术要求", "type": "textarea", "layout": "block"},
                {"key": "business_spec", "label": "商务要求", "type": "textarea", "layout": "block"},
            ]},
            {"key": "eval", "title": "第四章 评审办法", "fields": [
                {"key": "eval_method", "label": "评审方法", "type": "select", "layout": "inline",
                 "options": ["综合评分法", "最低评标价法", "性价比法"]},
                {"key": "eval_factors", "label": "评分因素及标准", "type": "textarea", "layout": "block"},
                {"key": "price_score", "label": "价格分计算方式", "type": "textarea", "layout": "block"},
            ]},
            {"key": "contract", "title": "第五章 合同主要条款", "fields": [
                {"key": "contract_terms", "label": "合同主要条款", "type": "textarea", "layout": "block"},
            ]},
        ],
    },
    "internal_demand": {
        "key": "internal_demand",
        "name": "内江市第一人民医院采购需求表",
        "subtitle": "（院内竞选 / 院内单一来源）",
        "sections": [
            {"key": "basic", "title": "第一部分　项目基本情况", "fields": [
                {"key": "demand_dept", "label": "需求科室", "type": "text", "layout": "inline"},
                {"key": "manage_dept", "label": "归口管理科室", "type": "text", "layout": "inline"},
                {"key": "project_name", "label": "项目名称", "type": "text", "layout": "inline", "required": True},
                {"key": "purchaser", "label": "采购单位", "type": "text", "layout": "inline"},
                {"key": "year", "label": "所属年度", "type": "text", "layout": "inline"},
                {"key": "compile_date", "label": "编制时间", "type": "date", "layout": "inline"},
                {"key": "category", "label": "品目类别", "type": "select", "layout": "inline",
                 "options": ["货物", "服务", "工程"]},
                {"key": "overview", "label": "项目概况", "type": "textarea", "layout": "block"},
                {"key": "has_consultant", "label": "是否聘请论证专家", "type": "select", "layout": "inline",
                 "options": ["否", "是"]},
                {"key": "consultant_note", "label": "论证情况说明", "type": "textarea", "layout": "block"},
            ]},
            {"key": "plan", "title": "第三部分　项目采购实施计划", "fields": [
                {"key": "proc_method", "label": "采购方式", "type": "select", "layout": "inline",
                 "options": ["院内竞选", "院内单一来源"]},
                {"key": "pkg_split", "label": "分包情况", "type": "select", "layout": "inline",
                 "options": ["不分包采购", "分包采购"]},
                {"key": "multi_year", "label": "是否跨年度采购", "type": "select", "layout": "inline",
                 "options": ["否", "是"]},
            ]},
            {"key": "target", "title": "第四部分　分包情况及标的情况", "fields": [
                {"key": "items", "label": "标的情况", "type": "table", "layout": "block", "columns": [
                    {"key": "pkg", "label": "包号"}, {"key": "code", "label": "编号"},
                    {"key": "name", "label": "标的名称"}, {"key": "price", "label": "单价(元)"},
                    {"key": "qty", "label": "数量"}, {"key": "unit", "label": "计量单位"},
                    {"key": "amount", "label": "标的金额(元)"}]},
                {"key": "budget", "label": "预算金额(元)", "type": "number", "layout": "inline"},
                {"key": "pkg_detail", "label": "具体分包说明", "type": "textarea", "layout": "block"},
            ]},
            {"key": "tech", "title": "第五部分　技术要求", "fields": [
                {"key": "tech_req", "label": "技术要求（标“★”为实质性条款）", "type": "textarea", "layout": "block"},
            ]},
            {"key": "biz", "title": "第六部分　商务要求", "fields": [
                {"key": "biz_req", "label": "商务要求", "type": "textarea", "layout": "block"},
            ]},
            {"key": "qualify", "title": "第七部分　资格要求", "fields": [
                {"key": "general_qual", "label": "一般资格要求", "type": "textarea", "layout": "block"},
                {"key": "special_qual", "label": "特殊资格要求", "type": "textarea", "layout": "block"},
            ]},
            {"key": "eval", "title": "第八部分　评审因素", "fields": [
                {"key": "eval_method", "label": "评审办法", "type": "select", "layout": "inline",
                 "options": ["综合评分法", "最低评标价法", "性价比法"]},
                {"key": "review_factors", "label": "评审因素", "type": "table", "layout": "block", "columns": [
                    {"key": "factor", "label": "评审因素"}, {"key": "score", "label": "评审分值"},
                    {"key": "objective", "label": "是否客观项"}, {"key": "standard", "label": "评审标准"}]},
            ]},
            {"key": "contract", "title": "第九部分　合同主要条款", "fields": [
                {"key": "contract_type", "label": "合同类型", "type": "text", "layout": "inline"},
                {"key": "contract_period", "label": "合同履行期限", "type": "text", "layout": "inline"},
                {"key": "contract_location", "label": "履约地点", "type": "text", "layout": "inline"},
                {"key": "payment_terms", "label": "支付约定", "type": "textarea", "layout": "block"},
                {"key": "acceptance_standard", "label": "验收交付标准和方法", "type": "textarea", "layout": "block"},
                {"key": "warranty", "label": "质量保修范围和保修期", "type": "textarea", "layout": "block"},
                {"key": "breach_dispute", "label": "违约责任与争议解决", "type": "textarea", "layout": "block"},
                {"key": "contract_other", "label": "合同其他条款", "type": "textarea", "layout": "block"},
            ]},
        ],
    },
}


def get_template(key):
    return TEMPLATES.get(key)


def all_fields(key):
    tpl = TEMPLATES.get(key) or {}
    return [f for s in tpl.get("sections", []) for f in s["fields"]]


# ── 从项目预填（首次创建时带出已知字段） ────────────────────────
def prefill_from_project(template_key, project):
    """project: Project ORM 对象。返回初始 data dict。"""
    if not project:
        return {}
    cat = ""
    line = (project.line or "")
    if line == "货物":
        cat = "货物"
    elif line in ("服务", "工程"):
        cat = line
    common = {
        "project_name": project.name or "",
        "project_no": project.number or "",
        "budget": project.amount if project.amount else "",
        "year": project.year or "",
        "purchaser": "内江市第一人民医院",
        "demand_dept": project.demand_dept or "",
        "procure_method": project.method or "",
        "category": cat,
    }
    if template_key == "procurement_demand":
        return common
    if template_key == "procurement_doc":
        return {
            "project_intro": f"{project.name or ''}，预算 {project.amount or ''} 元。",
        }
    if template_key == "internal_demand":
        return {
            "project_name": project.name or "",
            "purchaser": "内江市第一人民医院",
            "year": project.year or "",
            "demand_dept": project.demand_dept or "",
            "manage_dept": project.manage_dept or "",
            "budget": project.amount if project.amount else "",
            "category": cat,
            "proc_method": project.method if (project.method or "") in ("院内竞选", "院内单一来源") else "院内竞选",
        }
    return {}


# ── 完成度统计 ──────────────────────────────────────────────────
def _is_filled(f, data):
    v = data.get(f["key"])
    if f.get("type") == "table":
        return isinstance(v, list) and len(v) > 0
    return bool(str(v if v is not None else "").strip())


def progress(template_key, data):
    fields = all_fields(template_key)
    total = len(fields)
    filled = sum(1 for f in fields if _is_filled(f, data))
    return {"total": total, "filled": filled}


# ── Word 导出 ───────────────────────────────────────────────────
def generate_word(template_key, data):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    tpl = TEMPLATES.get(template_key) or {}
    doc = Document()
    # 标题
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(tpl.get("name", ""))
    r.bold = True
    r.font.size = Pt(18)
    if tpl.get("subtitle"):
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run(tpl["subtitle"])
        sr.font.size = Pt(10)

    for sec in tpl.get("sections", []):
        sh = doc.add_paragraph()
        shr = sh.add_run(sec["title"])
        shr.bold = True
        shr.font.size = Pt(13)
        for f in sec["fields"]:
            if f.get("type") == "table":
                lp = doc.add_paragraph()
                lp.add_run(f["label"] + "：").bold = True
                cols = f.get("columns", [])
                rows = data.get(f["key"]) or []
                if cols and rows:
                    tb = doc.add_table(rows=1, cols=len(cols))
                    tb.style = "Table Grid"
                    for i, c in enumerate(cols):
                        tb.rows[0].cells[i].text = c["label"]
                    for r in rows:
                        cells = tb.add_row().cells
                        for i, c in enumerate(cols):
                            cells[i].text = str(r.get(c["key"], "") or "")
                else:
                    doc.add_paragraph("（未填写）")
                continue
            val = str(data.get(f["key"], "") or "").strip()
            if f.get("layout") == "block":
                lp = doc.add_paragraph()
                lp.add_run(f["label"] + "：").bold = True
                doc.add_paragraph(val or "（未填写）")
            else:
                p = doc.add_paragraph()
                p.add_run(f["label"] + "：").bold = True
                p.add_run(val or "（未填写）")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
