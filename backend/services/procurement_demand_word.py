import io
import os
import json
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

TEMPLATE = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "..", "..", "医院模板", "文件汇总", "2.2内江市第一人民医院采购需求表.docx"
)

def _set_cell(table, row, col, text):
    """安全设置单元格文本（保留原有格式，只替换第一个段落的第一个run的文本）"""
    try:
        cell = table.cell(row, col)
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].text = text
            for r in cell.paragraphs[0].runs[1:]:
                r.text = ""
        elif cell.paragraphs:
            cell.paragraphs[0].text = text
        else:
            cell.text = text
    except Exception:
        pass

def _set_cell_text(table, row, col, text):
    """完全替换单元格文本"""
    try:
        cell = table.cell(row, col)
        cell.text = text
    except Exception:
        pass

def generate(demand, project):
    doc = Document(TEMPLATE)
    t = doc.tables[0]

    pname = demand.project_name or (project.name if project else "")
    year = demand.year or (str(project.year) if project and project.year else "")
    budget = f"{demand.budget_amount:.2f}" if demand.budget_amount else "XXXXXXXXX"

    # ── 表头 ─────────────────────────────────────────────────────
    # Row0: 需求科室
    _set_cell_text(t, 0, 0, f"需求科室\n{demand.demand_dept or ''}")
    _set_cell_text(t, 0, 11, f"归口管理科室\n{demand.manage_dept or '医学装备部'}")
    # Row1: 项目名称
    _set_cell_text(t, 1, 0, f"项目名称\n{pname}")

    # ── 第一部分 ──────────────────────────────────────────────────
    category_str = f"{'☑' if demand.category=='货物' else '□'}货物{'☑' if demand.category=='服务' else '□'}服务{'☑' if demand.category=='工程' else '□'}工程"
    part1 = (
        f"1.1采购单位：内江市第一人民医院\n"
        f"1.2所属年度：{year}年\n"
        f"1.3编制单位：内江市第一人民医院\n"
        f"1.4编制时间：{demand.compile_date or ''}\n"
        f"1.5项目所属分类：{category_str}\n"
        f"1.6预算金额（元）：{budget}\n"
        f"1.7项目概况：{demand.project_overview or 'XXXXXXXXX'}\n"
        f"1.8本项目是否有为采购项目提供整体设计、规范编制或者项目管理、监理、检测等服务的供应商：{demand.has_related_supplier or '否'}"
    )
    _set_cell_text(t, 3, 0, part1)

    # ── 第三部分 ──────────────────────────────────────────────────
    method = demand.procurement_method or "院内竞选"
    method_str = (
        f"{'☑' if method=='院内竞选' else '□'}院内竞选"
        f"{'☑' if method=='院内单一来源' else '□'}院内单一来源\n"
        f"{'☑' if method=='院内询价' else '□'}院内询价"
        f"{'☑' if method=='院内议价' else '□'}院内议价"
    )
    multi_year = f"{'☑' if demand.is_multi_year=='是' else '□'}是{'☑' if demand.is_multi_year=='否' else '□'}否"
    pkg_split = f"{'☑' if demand.package_split=='不分包采购' else '□'}不分包采购  {'☑' if demand.package_split=='分包采购' else '□'}分包采购"
    part3 = (
        f"3.1采购组织形式：□分散采购□政府集中采购☑自行采购\n"
        f"3.2采购方式：\n"
        f"3.2.3采购方式（自行采购）：{method_str}\n"
        f"3.3本项目是否单位自行组织采购：☑是□否\n"
        f"3.4采购包划分：{pkg_split}\n"
        f"3.5是否属于一签多年项目：{multi_year}"
    )
    _set_cell_text(t, 7, 0, part3)

    # ── 第四部分：标的 ─────────────────────────────────────────────
    try:
        items = json.loads(demand.items_json or "[]")
    except Exception:
        items = []
    if items:
        _set_cell_text(t, 11, 0, "1")
        _set_cell_text(t, 11, 1, items[0].get("item_no", "1-1"))
        _set_cell_text(t, 11, 3, items[0].get("category", ""))
        _set_cell_text(t, 11, 4, items[0].get("name", ""))
        _set_cell_text(t, 11, 7, str(items[0].get("unit_price", "")))
        _set_cell_text(t, 11, 8, str(items[0].get("quantity", "")))
        _set_cell_text(t, 11, 10, items[0].get("unit", ""))
        _set_cell_text(t, 11, 11, str(items[0].get("amount", "")))
        _set_cell_text(t, 11, 16, items[0].get("requirements", ""))

    # ── 第五部分：技术要求 ────────────────────────────────────────
    if demand.tech_requirements:
        _set_cell_text(t, 18, 0, f"一、包一\n{demand.tech_requirements}")

    # ── 第六部分：商务要求 ────────────────────────────────────────
    if demand.business_requirements:
        _set_cell_text(t, 19, 0, f"一、包一：\n{demand.business_requirements}")

    # ── 第九部分 ──────────────────────────────────────────────────
    actual = f"{'☑' if demand.contract_is_actual=='是' else '□'}是{'☑' if demand.contract_is_actual=='否' else '□'}否"
    part9 = (
        f"9.2是否为据实结算：{actual}\n"
        f"9.3合同履行期限：{demand.contract_period or ''}\n"
        f"9.5合同支付约定：{demand.payment_terms or ''}\n"
        f"9.10违约责任与解决争议的方法：{demand.breach_terms or ''}"
    )
    _set_cell_text(t, 30, 0, part9)

    # ── 第十部分 ──────────────────────────────────────────────────
    part10 = (
        f"10.6履约验收程序：{demand.acceptance_procedure or ''}\n"
        f"10.7履约验收时间：{demand.acceptance_time or ''}\n"
        f"10.9技术履约验收内容：{demand.acceptance_tech or ''}\n"
        f"10.10商务履约验收内容：{demand.acceptance_biz or ''}\n"
        f"10.11履约收标准：{demand.acceptance_standard or ''}"
    )
    _set_cell_text(t, 32, 0, part10)

    # ── 审签 ─────────────────────────────────────────────────────
    _set_cell_text(t, 35, 0, f"归口管理科室\n经办人：{demand.handler_name or ''}")
    _set_cell_text(t, 35, 11, f"归口管理科室\n负责人意见：{demand.manager_opinion or ''}")
    _set_cell_text(t, 36, 0, f"法律顾问意见\n{demand.legal_opinion or ''}")
    _set_cell_text(t, 37, 0, f"审计科意见\n{demand.audit_opinion or ''}")
    _set_cell_text(t, 38, 0, f"归口管理科室\n分管院领导意见：{demand.leader_opinion or ''}")
    _set_cell_text(t, 39, 0, f"总会计师意见\n{demand.cfo_opinion or ''}")
    _set_cell_text(t, 40, 0, f"院长意见\n{demand.president_opinion or ''}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"采购需求表_{project.number if project else 'unknown'}_{demand.status}.docx"
    return buf, filename
