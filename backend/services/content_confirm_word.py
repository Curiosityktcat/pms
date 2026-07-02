import io
import os
from docx import Document

TEMPLATE = os.path.abspath(os.path.join(
    os.path.dirname(__file__), "..", "..",
    "医院模板", "文件汇总", "1.（打印1份盖章）《院内竞选文件》内容确认表（第xx次）.docx"
))


def _set_cell_text(table, row, col, text):
    """完全替换单元格文本（保留单元格本身格式）。"""
    try:
        table.cell(row, col).text = text
    except Exception:
        pass


def _category_str(category):
    cats = ["货物", "服务", "工程"]
    return "".join(("☑" if category == c else "□") + c for c in cats)


def _hash_block(file_hashes):
    """把 [(文件名, sha256), ...] 拼成多行文本；单文件只填哈希值。"""
    pairs = [(n, h) for n, h in (file_hashes or []) if h]
    if not pairs:
        return ""
    if len(pairs) == 1:
        return pairs[0][1]
    return "\n".join(f"{n}：{h}" for n, h in pairs)


def _contact_str(contact, phone):
    """把联系人/联系方式拼成单元格文本。"""
    contact = (contact or "").strip()
    phone = (phone or "").strip()
    if contact and phone:
        return f"{contact}  {phone}"
    return contact or phone


def _set_para_text(para, text):
    """替换段落文本，保留首个 run 的格式。"""
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)


def generate(project, agency_name, *, file_hashes=None,
             contact_person="", contact_phone="", confirm_date=""):
    """按模板生成《院内竞选文件》内容确认表，返回 (BytesIO, filename)。

    填充：项目名称、采购编号、项目内容(分类勾选)、代理机构、采购方式、
    竞选文件 SHA256 哈希值、代理机构联系人及联系方式。
    其余（采购人审核结论/审核日期/盖章）留空，人工填写。
    """
    doc = Document(TEMPLATE)
    t = doc.tables[0]

    pname = (project.name or "").strip()
    pnumber = (project.number or "").strip()

    _set_cell_text(t, 0, 1, pname)                              # 项目名称（合并单元格）
    _set_cell_text(t, 1, 1, pnumber)                            # 采购编号
    _set_cell_text(t, 1, 3, _category_str(project.category))    # 项目内容（货物/服务/工程）
    _set_cell_text(t, 2, 1, (agency_name or "").strip())        # 代理机构
    _set_cell_text(t, 2, 3, _contact_str(contact_person, contact_phone))  # 代理机构联系人及联系方式
    _set_cell_text(t, 3, 3, (project.officer or "").strip())    # 采购人联系人（直接取项目经办人）
    _set_cell_text(t, 4, 1, _hash_block(file_hashes))           # 竞选文件哈希值（SHA256）
    _set_cell_text(t, 4, 3, (project.method or "院内竞选").strip())  # 采购方式

    # 审核确认日期 = 经办人确认采购文件当天（行5 审核结论段内，仅替换日期占位，保留正文）
    if confirm_date:
        for para in t.cell(5, 0).paragraphs:
            if "审核确认日期" in para.text:
                _set_para_text(para, f"审核确认日期  ：{confirm_date}")
                break

    buf = io.BytesIO()
    from services.docx_utils import strip_highlight
    strip_highlight(doc)                 # 清除模板黄色高亮占位印记
    doc.save(buf)
    buf.seek(0)
    filename = f"内容确认表_{pnumber or pname or '确认表'}.docx"
    return buf, filename
