# -*- coding: utf-8 -*-
"""采购文件 AI 定稿生成：初稿/模板(docx|doc) + 采购需求(docx|doc) → DeepSeek 段落级修订 → 定稿 docx。

设计（与 2026-07-02 项目32原型一致）：
- 保留初稿全部版式：模型只输出「哪些段落改成什么」，python-docx 原位替换，
  替换文本沿用原段落样式；多行文本拆段插入；连续「a | b | c」行转真表格。
- 模型与 key 取自 sys_config（scraper_model_api/scraper_api_key，与全站 AI 一致）。
"""
import copy as _copy
import json
import os
import re
import subprocess
import tempfile

from docx import Document
from docx.text.paragraph import Paragraph

SYSTEM = """你是政府采购代理机构的资深文件编制专家。任务：依据【采购需求】修订【采购文件初稿】，产出定稿修订清单。

要求：
1. 逐项核对初稿与采购需求：项目名称、预算/最高限价、数量、技术参数、商务要求（配送/结算/服务期）、资格条件、评审因素等，凡与需求不一致或缺失的，给出修订。
2. 只修改必须修改的段落，其他段落一律不动；修订后的文本必须是完整的段落全文（不是差异片段）。
3. 不得虚构需求里没有的内容；需求里明确的数字（限价、数量、日期）必须一字不差。
4. 修订文本可以多行（用 \\n 分行）：首行替换原段落，后续各行作为新段落插入其后；
   需要表格时用多行「列1 | 列2 | 列3」的形式表示（每行一条，会转成真正的表格）。
5. 初稿段落以 [编号] 开头给出。输出严格 JSON：
{"edits":[{"idx":编号,"text":"该段落修订后的完整文本","reason":"一句话原因"}],
 "summary":"总体修订说明（3~6条要点）"}
段落编号必须来自初稿；没有需要修订的地方就输出空 edits。"""


def _to_docx(path: str) -> str:
    """.doc（或其他 soffice 认识的格式）转临时 docx；docx 原样返回。"""
    if path.lower().endswith(".docx"):
        return path
    tmpdir = tempfile.mkdtemp(prefix="aidoc_")
    r = subprocess.run(
        ["soffice", "--headless", "--convert-to", "docx", "--outdir", tmpdir, path],
        capture_output=True, timeout=180)
    base = os.path.splitext(os.path.basename(path))[0] + ".docx"
    out = os.path.join(tmpdir, base)
    if r.returncode != 0 or not os.path.exists(out):
        raise RuntimeError("文件格式转换失败（仅支持 doc/docx）")
    return out


def _doc_text(path: str) -> str:
    d = Document(_to_docx(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        seen = set()
        for r in t.rows:
            for c in r.cells:
                k = id(c._tc)
                if k in seen:
                    continue
                seen.add(k)
                s = c.text.strip()
                if s:
                    parts.append(s)
    return "\n".join(parts)


def _set_para_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def _apply_edit(doc, idx, new_text):
    p = doc.paragraphs[idx]
    lines = [l.rstrip() for l in new_text.split("\n") if l.strip()]
    _set_para_text(p, lines[0])
    anchor = p._p
    rest = lines[1:]
    i = 0
    while i < len(rest):
        if " | " in rest[i] and i + 1 < len(rest) and " | " in rest[i + 1]:
            grp = []
            while i < len(rest) and " | " in rest[i]:
                grp.append([c.strip() for c in rest[i].split("|")])
                i += 1
            cols = max(len(r) for r in grp)
            tbl = doc.add_table(rows=len(grp), cols=cols)
            tbl.style = "Table Grid"
            for ri, row in enumerate(grp):
                for ci, val in enumerate(row):
                    if ci < cols:
                        tbl.rows[ri].cells[ci].text = val
            anchor.addnext(tbl._tbl)
            anchor = tbl._tbl
        else:
            np_el = _copy.deepcopy(p._p)
            np = Paragraph(np_el, p._parent)
            _set_para_text(np, rest[i])
            anchor.addnext(np_el)
            anchor = np_el
            i += 1


def generate_final_doc(draft_path: str, demand_path: str, out_path: str):
    """生成定稿。返回 (summary, applied_edits, usage_dict)。抛异常=失败。"""
    from models import db
    from models.sys_config import SysConfig

    demand = _doc_text(demand_path)
    doc = Document(_to_docx(draft_path))
    lines = [f"[{i}] {p.text.strip()}"
             for i, p in enumerate(doc.paragraphs) if p.text.strip()]
    draft_txt = "\n".join(lines)

    tbl_parts = []
    for ti, t in enumerate(doc.tables):
        seen, cells = set(), []
        for r in t.rows:
            for c in r.cells:
                k = id(c._tc)
                if k in seen:
                    continue
                seen.add(k)
                s = c.text.strip()
                if s:
                    cells.append(s.replace("\n", " ")[:120])
        tbl_parts.append(f"《表{ti + 1}》" + " | ".join(cells)[:1500])

    user_msg = (f"【采购需求（全文）】\n{demand}\n\n"
                f"【采购文件初稿（带段落编号；表格内容附后仅供核对，不可修订表格）】\n{draft_txt}\n\n"
                f"【初稿内表格内容（只读）】\n" + "\n".join(tbl_parts))

    key = db.session.get(SysConfig, "scraper_api_key")
    model = db.session.get(SysConfig, "scraper_model_name")
    if not key or not key.value:
        raise RuntimeError("未配置 AI 模型 API Key（后台管理→大模型配置）")
    from openai import OpenAI
    client = OpenAI(api_key=key.value, base_url="https://api.deepseek.com")
    resp = client.chat.completions.create(
        model=(model.value if model and model.value else "deepseek-v4-flash"),
        messages=[{"role": "system", "content": SYSTEM},
                  {"role": "user", "content": user_msg}],
        response_format={"type": "json_object"},
        temperature=0,
        max_tokens=8000,
    )
    choice = resp.choices[0]
    if choice.finish_reason == "length":
        raise RuntimeError("模型输出被截断，文件过大，请人工编制")
    out = json.loads(choice.message.content)
    edits = out.get("edits") or []

    applied = []
    for e in sorted(edits, key=lambda x: -int(x["idx"])):
        idx = int(e["idx"])
        if not (0 <= idx < len(doc.paragraphs)):
            continue
        old = doc.paragraphs[idx].text
        new = str(e.get("text") or "").strip()
        if not new or new == old.strip():
            continue
        _apply_edit(doc, idx, new)
        applied.append({"idx": idx, "reason": e.get("reason", ""),
                        "old": old[:60], "new": new.split("\n")[0][:60]})
    applied.reverse()

    doc.save(out_path)
    usage = {
        "prompt_tokens": resp.usage.prompt_tokens,
        "completion_tokens": resp.usage.completion_tokens,
        "total_tokens": resp.usage.total_tokens,
    }
    return out.get("summary", ""), applied, usage
