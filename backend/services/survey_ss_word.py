"""调研公告（6.2）与单一来源公示（6.4）的 Word 生成。

体例来源：
  · 调研公告 —— 照内江一院官网已发布的市场调研公告（云算力 id/7075、
    法律顾问 id/7078 两份实样）：项目概况 → 调研内容及要求 → 供应商资格要求
    → 报价要求 → 提交资料及方式 → 联系方式 → 特别说明。
    结尾必须有「本次调研结果与采购结果无必然联系」这句，两份实样都有。
  · 单一来源公示 —— 官网没有先例可抄，按《政府采购非招标采购方式管理办法》
    （财政部令第74号）第38条的法定必备内容组织：采购人与项目名称、拟采购
    货物或服务说明、采用单一来源的原因、拟定唯一供应商名称地址、专业人员
    论证意见（含姓名/工作单位/职称）、公示期、异议接收部门与联系方式。
"""
import io
import json
import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PURCHASER = "内江市第一人民医院"
PURCHASER_ADDR = "四川省内江市市中区沱中路41号、汉安大道西段1866号"
PURCHASER_ZIP = "641000"

SURVEY_NOTE_DEFAULT = (
    "1.本次调研仅用于采购需求论证，医院不保证采纳任何单位提供的方案。\n"
    "2.参与单位应保证所提交资料真实有效，弄虚作假的取消其参与资格。\n"
    "3.本次调研结果与本项目的采购结果无任何必然联系。\n"
    "4.本公告自发布之日起生效，医院保留对本公告的最终解释权。"
)


def _para(doc, text, size=14, bold=False, center=False, indent=True,
          font="仿宋_GB2312"):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    return p


def _title(doc, text):
    _para(doc, text, size=18, bold=True, center=True, indent=False,
          font="方正小标宋简体")
    doc.add_paragraph()


def _section(doc, text):
    _para(doc, text, size=14, bold=True, indent=False, font="黑体")


def _multiline(doc, text):
    """把多行文本逐行输出，保留用户排的条目顺序。"""
    for line in (text or "").splitlines():
        line = line.strip()
        if line:
            _para(doc, line)


def _tail_contact(doc, ann, project):
    _section(doc, "联系方式")
    _para(doc, f"采购人：{PURCHASER}")
    _para(doc, f"地址：{PURCHASER_ADDR}")
    _para(doc, f"邮编：{PURCHASER_ZIP}")
    if ann.agency_contact:
        _para(doc, f"联系人：{ann.agency_contact}")
    if ann.agency_contact_phone:
        _para(doc, f"联系电话：{ann.agency_contact_phone}")
    if ann.agency_email:
        _para(doc, f"电子邮箱：{ann.agency_email}")


def _sign(doc):
    doc.add_paragraph()
    now = datetime.datetime.now()
    _para(doc, PURCHASER, center=True, indent=False)
    _para(doc, f"{now.year}年{now.month}月{now.day}日", center=True, indent=False)


def build_survey(ann, project):
    """调研公告（市场调研公告）。返回 (BytesIO, 文件名)。"""
    doc = Document()
    name = project.name if project else ""
    _title(doc, f"{name}市场调研公告")

    _section(doc, "一、项目概况")
    _para(doc, f"项目名称：{name}")
    if project and project.number:
        _para(doc, f"项目编号：{project.number}")
    if ann.project_intro:
        _multiline(doc, ann.project_intro)

    _section(doc, "二、调研内容及要求")
    _multiline(doc, ann.survey_content)

    if (ann.survey_qualification or "").strip():
        _section(doc, "三、参与单位资格要求")
        _multiline(doc, ann.survey_qualification)

    if (ann.survey_quote_req or "").strip():
        _section(doc, "四、报价要求")
        _multiline(doc, ann.survey_quote_req)

    _section(doc, "五、需提交的资料及方式")
    if (ann.survey_materials or "").strip():
        _multiline(doc, ann.survey_materials)
    if ann.survey_deadline:
        _para(doc, f"提交截止时间：{ann.survey_deadline}")
    if (ann.survey_submit_way or "").strip():
        _multiline(doc, ann.survey_submit_way)

    _tail_contact(doc, ann, project)

    _section(doc, "特别说明")
    _multiline(doc, ann.survey_note or SURVEY_NOTE_DEFAULT)

    _sign(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf, f"市场调研公告_{(project.number if project else '') or name}.docx"


def build_single_source(ann, project):
    """单一来源采购公示。返回 (BytesIO, 文件名)。"""
    doc = Document()
    name = project.name if project else ""
    _title(doc, f"{name}单一来源采购公示")

    _section(doc, "一、项目基本情况")
    _para(doc, f"采购人：{PURCHASER}")
    _para(doc, f"项目名称：{name}")
    if project and project.number:
        _para(doc, f"项目编号：{project.number}")
    if project and project.amount:
        _para(doc, f"采购预算：{project.amount} 元")

    _section(doc, "二、拟采购的货物或服务说明")
    _multiline(doc, ann.ss_goods_desc)

    _section(doc, "三、采用单一来源采购方式的原因及相关说明")
    _multiline(doc, ann.ss_reason)

    _section(doc, "四、拟定的唯一供应商")
    _para(doc, f"名称：{ann.ss_supplier_name or ''}")
    _para(doc, f"地址：{ann.ss_supplier_addr or ''}")

    _section(doc, "五、专业人员论证意见")
    try:
        experts = json.loads(ann.ss_experts_json or "[]")
    except Exception:
        experts = []
    if experts:
        for i, e in enumerate(experts, 1):
            who = "；".join(x for x in [
                f"姓名：{e.get('name', '')}",
                f"工作单位：{e.get('org', '')}",
                f"职称：{e.get('title', '')}",
            ] if x.split("：", 1)[1])
            _para(doc, f"{i}. {who}")
            if e.get("opinion"):
                _para(doc, f"论证意见：{e['opinion']}")
    else:
        _para(doc, "（待补充专业人员论证意见）")

    _section(doc, "六、公示期限")
    if ann.ss_publicity_start and ann.ss_publicity_end:
        _para(doc, f"{ann.ss_publicity_start} 至 {ann.ss_publicity_end}"
                   f"（不少于 5 个工作日）")
    else:
        _para(doc, "自本公示发布之日起 5 个工作日")

    _section(doc, "七、异议的接收")
    _para(doc, "任何供应商、单位或者个人对采用单一来源采购方式公示有异议的，"
               "可以在公示期内将书面意见反馈至下列部门。")
    if ann.ss_objection_dept:
        _para(doc, f"接收部门：{ann.ss_objection_dept}")
    if ann.ss_objection_contact:
        _para(doc, f"联系人：{ann.ss_objection_contact}")
    if ann.ss_objection_phone:
        _para(doc, f"联系电话：{ann.ss_objection_phone}")
    _para(doc, f"地址：{ann.ss_objection_addr or PURCHASER_ADDR}")

    _tail_contact(doc, ann, project)
    _sign(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf, f"单一来源采购公示_{(project.number if project else '') or name}.docx"
