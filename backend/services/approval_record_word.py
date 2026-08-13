"""《审批过程记录表》Word 生成 —— 归档件之一。

采购文件这类环节可能来回驳回多次，光看最终稿看不出中间改了什么、
为什么改。归档时把 approval_logs 里该项目的全部往返按时间顺序拍成一张表，
连同其他要件一起放进归档文件夹，日后追溯有据可查。
"""
import io
import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from services.approval_log import render_rows

COLS = ["序号", "时间", "环节", "轮次", "动作", "操作人", "原因/说明"]
WIDTHS = [0.8, 3.2, 3.4, 2.0, 2.4, 1.8, 6.0]   # 厘米，合计约 19.6


def _set_font(run, size=10.5, bold=False, font="仿宋_GB2312"):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def build(project):
    """返回 (BytesIO, 文件名)。无往返记录时也出表，注明"无驳回记录"。"""
    rows = render_rows(project.id)

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Pt(48)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(title.add_run("审批过程记录表"), size=18, bold=True, font="方正小标宋简体")

    meta = doc.add_paragraph()
    _set_font(meta.add_run(
        f"项目名称：{project.name or ''}\n"
        f"项目编号：{project.number or ''}\n"
        f"采购方式：{project.method or ''}　　经办人：{project.officer or ''}\n"
        f"生成时间：{datetime.datetime.now().strftime('%Y年%m月%d日 %H:%M')}"
    ), size=10.5)

    table = doc.add_table(rows=1, cols=len(COLS))
    table.style = "Table Grid"
    for i, name in enumerate(COLS):
        cell = table.rows[0].cells[i]
        cell.text = ""
        _set_font(cell.paragraphs[0].add_run(name), size=10, bold=True, font="黑体")

    if rows:
        for r in rows:
            cells = table.add_row().cells
            for i, key in enumerate(COLS):
                cells[i].text = ""
                _set_font(cells[i].paragraphs[0].add_run(str(r.get(key, ""))), size=9.5)
    else:
        cells = table.add_row().cells
        cells[0].merge(cells[len(COLS) - 1])
        _set_font(cells[0].paragraphs[0].add_run("本项目无审批往返记录（各环节均一次通过）"), size=9.5)

    from docx.shared import Cm
    for row in table.rows:
        for i, w in enumerate(WIDTHS):
            row.cells[i].width = Cm(w)

    tail = doc.add_paragraph()
    rejects = sum(1 for r in rows if r.get("动作") in ("驳回", "不确认采购结果"))
    _set_font(tail.add_run(
        f"\n本项目共发生审批往返 {len(rows)} 次，其中驳回/不确认 {rejects} 次。"
        "本表由系统按操作留痕自动生成，不可手工修改。"
    ), size=10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    name = f"审批过程记录表_{project.number or project.name or project.id}.docx"
    return buf, name
