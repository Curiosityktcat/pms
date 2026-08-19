# -*- coding: utf-8 -*-
"""富文本字段：段落 + 表格，出稿时变成 Word 里真正的段落和表格。

黄新博 2026-08-19 ⑦：「技术要求……还需要可以直接放表格进去，
因为很多时候需要用表格去表达清楚采购需求；同理，商务需求也是这样。」

不用自由 HTML，改成**结构化的块**：
    [{"kind": "p", "text": "..."},
     {"kind": "table", "header": ["参数","要求"], "rows": [["屏幕","≥12吋"], ...]}]
理由：HTML 到 Word 的转换永远有对不上的地方（样式、嵌套、脏标签），
而这是要盖章对外发的文件。结构化的块少而明确，渲染结果可预期。
纯文字的老数据（就是一个字符串）也认，按段落切开——不用迁移。
"""
import json


def parse(value):
    """把字段值解析成块列表。老数据是纯字符串，按换行切成段落。"""
    if value is None or value == "":
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, str):
        v = value.strip()
        if v.startswith("[") or v.startswith("{"):
            try:
                data = json.loads(v)
            except Exception:                                # noqa: BLE001
                data = None
            if isinstance(data, list):
                return data
            if isinstance(data, dict) and isinstance(data.get("blocks"), list):
                return data["blocks"]
        return [{"kind": "p", "text": ln} for ln in v.splitlines() if ln.strip()]
    return []


def to_plain(value):
    """退化成纯文本（给完成度统计、搜索这类地方用）。"""
    out = []
    for b in parse(value):
        if b.get("kind") == "table":
            for row in [b.get("header") or []] + (b.get("rows") or []):
                out.append("　".join(str(c) for c in row))
        else:
            out.append(str(b.get("text") or ""))
    return "\n".join(x for x in out if x)


def _set_borders(table):
    """直接画边框，不用 "Table Grid" 样式名。

    模板文档里不一定有这个内置样式——《2.2采购需求表》就没有，
    照搬会抛 KeyError: no style with name 'Table Grid'（实测撞到）。
    换成写 tblBorders，跟模板带什么样式无关。
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def to_subdoc(tpl, value):
    """渲染成 docxtpl 的 Subdoc——Word 里真正的段落和表格。

    tpl 是 DocxTemplate 实例。返回 None 表示这个字段是空的，
    调用方直接填空字符串即可。
    """
    blocks = parse(value)
    if not blocks:
        return None
    sd = tpl.new_subdoc()
    for b in blocks:
        if b.get("kind") == "table":
            header = b.get("header") or []
            rows = b.get("rows") or []
            ncol = max([len(header)] + [len(r) for r in rows] or [1]) or 1
            t = sd.add_table(rows=0, cols=ncol)
            _set_borders(t)
            if header:
                cells = t.add_row().cells
                for i in range(ncol):
                    txt = str(header[i]) if i < len(header) else ""
                    cells[i].text = ""
                    run = cells[i].paragraphs[0].add_run(txt)
                    run.bold = True
            for r in rows:
                cells = t.add_row().cells
                for i in range(ncol):
                    cells[i].text = str(r[i]) if i < len(r) else ""
        else:
            txt = str(b.get("text") or "")
            if txt.strip():
                sd.add_paragraph(txt)
    return sd
