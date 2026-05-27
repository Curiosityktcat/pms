"""
授权函生成服务
使用 python-docx 填充模板，输出 .docx 文件
"""
import os, re, tempfile
from copy import deepcopy
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
PMS_ROOT = os.path.abspath(os.path.join(HERE, '..', '..'))
TEMPLATE_PATH = os.path.join(
    PMS_ROOT,
    '医院模板', '文件汇总',
    '3.（打印1份盖章）内江市第一人民医院授权函.docx'
)

CN_NUMS = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']


def _parse_bid_date(bid_time: str):
    """从 bid_time 字符串提取年/月/日，解析失败则原文作为年字段。"""
    if not bid_time:
        return '', '', ''
    m = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', bid_time)
    if m:
        return m.group(1) + '年', m.group(2), m.group(3)
    m = re.search(r'(\d{4})[/-](\d{1,2})[/-](\d{1,2})', bid_time)
    if m:
        return m.group(1) + '年', str(int(m.group(2))), str(int(m.group(3)))
    return bid_time, '', ''


def get_round_suffix(round_number: int) -> str:
    if round_number <= 1:
        return ''
    idx = round_number - 1
    cn = CN_NUMS[idx] if idx < len(CN_NUMS) else str(round_number)
    return f'（第{cn}次）'


def _set_rep_runs(para_elem, name: str, id_card: str):
    """在段落 XML 元素中设置代表的姓名（run[2]）和身份证号（run[3]）。"""
    w_runs = para_elem.findall('.//' + qn('w:r'))
    if len(w_runs) > 3:
        t_name = w_runs[2].find(qn('w:t'))
        t_card = w_runs[3].find(qn('w:t'))
        if t_name is not None:
            t_name.text = name
            # 保留空白字符
            t_name.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        if t_card is not None:
            t_card.text = f'（身份证号：{id_card or ""}）'
            t_card.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def generate(project, supervisor, representatives, agency_name: str,
             round_number: int = 1, bid_time_override: str = '') -> str:
    """
    生成授权函 Word 文件，返回临时文件路径（调用方负责删除）。

    project:         Project ORM 对象
    supervisor:      People ORM 对象（监督人员，单人）
    representatives: People ORM 对象列表（采购人代表，支持多人）
    agency_name:     代理机构全称
    round_number:    开标次数（1=第一次，2=第二次……）
    bid_time_override: 覆盖开标时间
    """
    if not representatives:
        raise ValueError('至少需要选择一位采购人代表')

    doc = Document(TEMPLATE_PATH)
    paras = doc.paragraphs
    project_name_in_letter = project.name + get_round_suffix(round_number)
    bid_time = bid_time_override.strip() if bid_time_override else (project.bid_time or '')

    # ── 段落 [1]：代理机构 + 项目名称 ─────────────────────────────────
    paras[1].runs[0].text = agency_name or '（代理机构）'
    paras[1].runs[4].text = project_name_in_letter

    # ── 段落 [2]：项目编号 ────────────────────────────────────────────
    paras[2].runs[1].text = project.number or ''
    paras[2].runs[2].text = ''

    # ── 段落 [3]：开标日期 ────────────────────────────────────────────
    year_str, month_str, day_str = _parse_bid_date(bid_time)
    paras[3].runs[1].text = year_str
    paras[3].runs[2].text = month_str
    paras[3].runs[3].text = '月' if month_str else ''
    paras[3].runs[4].text = day_str
    paras[3].runs[5].text = '日' if day_str else ''

    # ── 段落 [4]：监督人员 ────────────────────────────────────────────
    paras[4].runs[1].text = supervisor.name
    paras[4].runs[2].text = f'（身份证号：{supervisor.id_card or ""}）'

    # ── 段落 [6]：采购人代表（支持多人）──────────────────────────────
    p6 = paras[6]
    # 保存原始模板（在修改之前深拷贝，用于后续多人复制）
    p6_template_elem = deepcopy(p6._element)

    # 填入第一位代表
    p6.runs[2].text = representatives[0].name
    p6.runs[3].text = f'（身份证号：{representatives[0].id_card or ""}）'

    # 第二位及以后：插入模板副本到 p6 之后
    last_rep_elem = p6._element
    for rep in representatives[1:]:
        new_elem = deepcopy(p6_template_elem)
        last_rep_elem.addnext(new_elem)
        last_rep_elem = new_elem
        _set_rep_runs(new_elem, rep.name, rep.id_card or '')

    # ── 第六部分：附身份证照片 ────────────────────────────────────────
    def add_photo(label: str, photo_rel: str):
        doc.add_paragraph(label)
        full = os.path.join(PMS_ROOT, photo_rel) if photo_rel else None
        if full and os.path.exists(full):
            p = doc.add_paragraph()
            try:
                p.add_run().add_picture(full, width=Inches(4.5))
            except Exception:
                doc.add_paragraph(f'（照片加载失败：{photo_rel}）')
        else:
            doc.add_paragraph('（未找到照片文件）')

    # 监督人照片
    add_photo(f'监督人员 {supervisor.name} 身份证：', supervisor.photo or '')

    # 所有代表照片
    for rep in representatives:
        doc.add_paragraph('')
        add_photo(f'采购人代表 {rep.name} 身份证：', rep.photo or '')

    # ── 保存临时文件 ──────────────────────────────────────────────────
    tmp = tempfile.NamedTemporaryFile(
        suffix='.docx', delete=False,
        prefix=f'授权函_{project.number or "draft"}_'
    )
    tmp.close()
    doc.save(tmp.name)
    return tmp.name
