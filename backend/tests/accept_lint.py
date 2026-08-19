"""模板体检验收：坏模板必须在**上传那一刻**就被抓出来。

pdt 第六节：「认不出的占位符要当场列给你看，而不是等生成出来一片空白才发现。」
每一条查的都是真踩过的坑。
"""
import os, sys, tempfile
sys.path.insert(0, ".")
ok, bad = 0, []


def check(t, c, e=""):
    global ok
    if c: ok += 1; print(f"OK   {t} {e}")
    else: bad.append(t); print(f"FAIL {t} {e}")


from docx import Document
from services import template_lint, field_dict

TMP = tempfile.mkdtemp(prefix="lint_")


def make(lines, tables=None):
    d = Document()
    for ln in lines:
        d.add_paragraph(ln)
    for rows in (tables or []):
        t = d.add_table(rows=0, cols=max(len(r) for r in rows))
        for r in rows:
            cs = t.add_row().cells
            for i, v in enumerate(r):
                cs[i].text = v
    p = os.path.join(TMP, f"t{len(os.listdir(TMP))}.docx")
    d.save(p)
    return p


# ═══ 现役模板应当零错 ═══════════════════════════════════════════
names = [f["name"] for f in field_dict.load()]
e, w, n = template_lint.lint(
    "/home/huangxb/pms/backend/templates_docx/2.2采购需求表.docx", names)
check("① 现役模板零错误", not e, f"{e[:2]}")
check("① 占位符认得出来", len(n) > 40, f"{len(n)} 个")

# ═══ ☑ ☐ ═══════════════════════════════════════════════════════
e, _, _ = template_lint.lint(make(["资格审查：☑ 甲方 ☐ 乙方"]))
check("② 用了 ☑☐ 会报错", any("字形" in x for x in e), f"{e[:1]}")
e, _, _ = template_lint.lint(make(["资格审查：■ 甲方 □ 乙方"]))
check("② 用 ■□ 不报错", not any("字形" in x for x in e))

# ═══ 循环标记 ═══════════════════════════════════════════════════
e, _, _ = template_lint.lint(make(["{%tr for r in 标的 %}", "{{ r.名称 }}"]))
check("③ for 少了 endfor 会报错", any("不成对" in x for x in e), f"{e[:1]}")
e, _, _ = template_lint.lint(make(["{%tr for r in 标的 %}{{ r.名称 }}{%tr endfor %}"]))
check("③ for 和 endfor 同一段会报错", any("同一段" in x for x in e), f"{e[:1]}")
e, _, _ = template_lint.lint(make(["{%tr for r in 标的 %}", "{{ r.名称 }}", "{%tr endfor %}"]))
check("③ 正确写法不报错", not e, f"{e[:1]}")

# ═══ 循环变量没有对应的 for ═════════════════════════════════════
e, _, _ = template_lint.lint(make(["{{ 包.预算金额 }}"]))
check("④ 用了循环变量却没 for 会报错",
      any("没有对应的" in x for x in e), f"{e[:1]}")

# ═══ 占位符没闭合 ═══════════════════════════════════════════════
e, _, _ = template_lint.lint(make(["预算：{{ 预算金额 "]))
check("⑤ 占位符没闭合会报错", any("没闭合" in x for x in e), f"{e[:1]}")

# ═══ 名字里有标点 ═══════════════════════════════════════════════
e, _, _ = template_lint.lint(make(["{{ 需求质疑、答复 }}"]))
check("⑥ 名字带标点会报错", any("标点" in x for x in e), f"{e[:1]}")

# ═══ 不在字典里的名字给提醒（不拦） ═════════════════════════════
e, w, _ = template_lint.lint(make(["{{ 我瞎编的字段 }}"]), ["项目名称"])
check("⑦ 字典里没有的只提醒不报错",
      not e and any("不在字段字典里" in x for x in w), f"错{len(e)} 提醒{len(w)}")

# ═══ 打不开的文件 ═══════════════════════════════════════════════
p = os.path.join(TMP, "bad.docx")
open(p, "wb").write(b"not a docx")
e, _, _ = template_lint.lint(p)
check("⑧ 打不开的文件给人话", any("打不开" in x for x in e), f"{e[:1]}")

print(f"\n通过 {ok} 项" + (f"，失败 {len(bad)}：{bad}" if bad else "，全部通过"))
sys.exit(1 if bad else 0)
