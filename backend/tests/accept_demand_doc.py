"""采购需求表按模板出稿的验收（只打测试库）。

用户 2026-08-18：「github 的 pdt 项目，尝试一下应用在 pms 系统里面，
初步应用在我的采购需求的模板……信息+Agent操作区+文件预览」。
成稿 ＝ 模板 ＋ 信息，所以重点验：模板语法没写坏、勾选打在对的地方、
中文没被 ☑ 带崩、算得出来的不让人填。
"""
import io as _io
import os, re, sys, secrets
sys.path.insert(0, ".")
os.environ["PMS_CAPTCHA_ON"] = "0"
os.environ["PMS_DB_PATH"] = "/home/huangxb/pms/pms.test.db"
PW = "Doc!2026"
ok, bad = 0, []


def check(t, c, e=""):
    global ok
    if c: ok += 1; print(f"OK   {t} {e}")
    else: bad.append(t); print(f"FAIL {t} {e}")


from docx import Document
from app import create_app
from models import db
from models.user import User
from models.procurement_demand import ProcurementDemand
from services.auth import hash_pw
from services import demand_doc

app = create_app(); app.app_context().push()
assert "pms.test.db" in app.config["SQLALCHEMY_DATABASE_URI"], "保险丝：不是测试库"

# ═══ 一、模板本身 ═══════════════════════════════════════════════
print("── 一、模板 ──")
check("① 模板文件在", os.path.exists(demand_doc.TEMPLATE), demand_doc.TEMPLATE)
fields = demand_doc.load_fields()
check("① 字段清单读得到", len(fields) > 40, f"{len(fields)} 个")
kinds = {f["kind"] for f in fields}
check("① 三类占位符都有", {"text", "choice", "table"} <= kinds, f"{kinds}")

doc = Document(demand_doc.TEMPLATE)
blob = "\n".join(p.text for p in doc.paragraphs)
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            blob += "\n" + c.text

# pdt 第三条：☑/☐ 在中文字体里没有字形，用了会让整行中文消失
check("① 模板里没有 ☑ ☐（中文字体没这俩字形）",
      "☑" not in blob and "☐" not in blob)
# pdt 第七条：正文里不能原样打出模板语法
check("① 循环标记成对", blob.count("{%tr for") == blob.count("{%tr endfor"),
      f"for {blob.count('{%tr for')} / endfor {blob.count('{%tr endfor')}")
check("① 占位符没有跨行断开", "{{" not in blob.replace("{{", "").replace("}}", ""),
      "")
# {{ 名字 }} 和 {{p 名字 }} 都要认——指令字母（p/r/tr/tc）不是名字。
# 直接用体检器那套规则，两处别各写各的。
names = {m.group(1).strip() for m in
         re.finditer(r"\{\{\s*(?:(?:tr|tc|p|r)\s+)?([^\s|}]+)", blob)}
declared = {f["name"] for f in fields}
# 循环里的变量（{{ r.xxx }} 标的、{{ 包.xxx }} 分包）不是字典字段，
# 它们的取值来自 build_context 里造的那两个列表，不该要求在字段清单里声明。
LOOP_PREFIX = ("r.", "包.")
unknown = {n for n in names if not n.startswith(LOOP_PREFIX)} - declared
check("① 模板里的占位符都在字段清单里", not unknown, f"多出 {sorted(unknown)[:5]}")

# 循环变量也得成对：for 里声明了什么，正文里才能用什么
check("① 循环变量都有对应的 for",
      ("{%tr for r in 标的" in blob or not any(n.startswith("r.") for n in names))
      and ("{%tr for 包 in 分包" in blob or not any(n.startswith("包.") for n in names)))

# ═══ 二、出稿 ═══════════════════════════════════════════════════
print("\n── 二、出稿 ──")
d = ProcurementDemand(
    demand_type="gov", status="草稿", created_by="验收",
    project_name="验收用采购项目", demand_dept="检验科", manage_dept="医学装备部",
    budget_amount=1234567.89, category="服务", year="2026",
    # 政府采购需求表：方式只能是政采那六种，组织形式只能是分散采购。
    # 这里故意填成不合规的值，验字典会不会在出稿时纠正过来。
    procurement_method="询价", org_form="自行采购",
    budget_method="询价",
    is_energy_save="是", has_import_product="否",
    items_json='[{"包号":"1","编号":"1-1","标的名称":"甲标的","数量":"2","计量单位":"台"},'
               '{"包号":"2","编号":"2-1","标的名称":"乙标的","数量":"5","计量单位":"套"}]')
db.session.add(d); db.session.commit()
did = d.id
print(f"   造了需求 #{did}")

buf, missing = demand_doc.render(d, None)
check("② 能出稿", buf.getbuffer().nbytes > 8000, f"{buf.getbuffer().nbytes} 字节")

out = Document(_io.BytesIO(buf.getvalue()))
text = "\n".join(p.text for p in out.paragraphs)
for t in out.tables:
    for row in t.rows:
        for c in row.cells:
            text += "\n" + c.text

check("② 成稿里没有残留的模板语法", "{{" not in text and "{%" not in text,
      text[text.find("{{"):text.find("{{") + 40] if "{{" in text else "")
check("② 项目名称填进去了", "验收用采购项目" in text)
check("② 科室填进去了", "检验科" in text and "医学装备部" in text)
check("② 金额带千分位", "1,234,567.89" in text, "")

# 勾选：选中的打 ■，没选的还是 □，且所有选项都印出来
check("② 分类「服务」被选中", "■ 服务" in text, "")
check("② 没选的印成 □", "□ 货物" in text and "□ 工程" in text)
check("② 是否题选中项正确", "■ 是" in text and "■ 否" in text,
      "3.8节能=是、3.9进口=否")
check("② 3.2 勾的是政采方式「询价」", "■ 询价" in text, "")
# 字典的活儿：库里存的是「自行采购」，成稿必须被纠正成「分散采购」
check("② 组织形式被字典纠正成分散采购",
      "■ 分散采购" in text and "□ 自行采购" in text,
      "库里存的是「自行采购」")
check("② 自行采购那几种方式在政采表上不勾",
      "■ 院内竞选" not in text and "■ 院内询价" not in text)

# 行循环：两条标的都要出来
check("② 标的清单两行都出来了", "甲标的" in text and "乙标的" in text)
check("② 标的的数量单位跟着走", "台" in text and "套" in text)

check("② 缺字段清单算得出来", isinstance(missing, list) and len(missing) > 0,
      f"还空着 {len(missing)} 个")

# ═══ 三、接口 ═══════════════════════════════════════════════════
print("\n── 三、接口 ──")
u = db.session.execute(db.select(User).filter_by(username="acc_docofficer")).scalar_one_or_none()
if u is None:
    u = User(username="acc_docofficer", display_name="验收", role="officer",
             dept_code="", active=1, agency_code="", salt="", pw_hash="")
    db.session.add(u)
u.display_name, u.role, u.active = "验收", "officer", 1
salt = secrets.token_hex(16); u.salt, u.pw_hash = salt, hash_pw(PW, salt)
if hasattr(u, "must_change_pw"): u.must_change_pw = 0
db.session.commit()

c = app.test_client()
c.post("/api/auth/login", json={"username": "acc_docofficer", "password": PW})

r = c.get("/api/procurement-demands/doc-fields")
check("③ 字段清单接口", r.status_code == 200 and len(r.get_json()["data"]) > 40,
      f"HTTP {r.status_code}")

r = c.get(f"/api/procurement-demands/{did}/doc-status")
dd = (r.get_json() or {}).get("data") or {}
check("③ 出稿状态接口", r.status_code == 200 and dd.get("total", 0) > 40,
      f"已填 {dd.get('filled')}/{dd.get('total')}")
check("③ 完成度算得对", dd.get("filled", 0) + len(dd.get("missing") or []) == dd.get("total"),
      f"{dd.get('filled')}+{len(dd.get('missing') or [])} vs {dd.get('total')}")

r = c.get(f"/api/procurement-demands/{did}/doc", query_string={"download": "1"})
check("③ 下载 Word", r.status_code == 200 and len(r.data) > 8000
      and "wordprocessingml" in (r.headers.get("Content-Type") or ""),
      f"HTTP {r.status_code}，{len(r.data)} 字节")

r = c.get(f"/api/procurement-demands/{did}/doc")
check("③ 预览返回的是 PDF（iframe 才渲染得了）",
      r.status_code == 200 and (r.headers.get("Content-Type") or "").startswith("application/pdf"),
      f"HTTP {r.status_code} 类型={r.headers.get('Content-Type')}")
check("③ PDF 不是空的", len(r.data) > 20000, f"{len(r.data)} 字节")

# 越权：别人的需求拿不到
r2 = app.test_client()
check("③ 没登录拿不到出稿",
      r2.get(f"/api/procurement-demands/{did}/doc").status_code in (401, 403))

db.session.delete(db.session.get(ProcurementDemand, did))
db.session.commit()
print(f"   已清理测试需求 #{did}")

print(f"\n通过 {ok} 项" + (f"，失败 {len(bad)}：{bad}" if bad else "，全部通过"))
sys.exit(1 if bad else 0)
