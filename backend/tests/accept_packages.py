"""分包验收：每个包一份第四~第八部分（黄新博 2026-08-19 ⑥⑪）。

「每个包就是一个独立的合同，具体需求与实施情况都不一样」——
所以重点验：包与包之间**互不串味**，成稿里各出一份完整的第四~第八部分。
"""
import io as _io
import json, os, secrets, sys
sys.path.insert(0, ".")
os.environ["PMS_CAPTCHA_ON"] = "0"
os.environ["PMS_DB_PATH"] = "/home/huangxb/pms/pms.test.db"
PW = "Pkg!2026"
ok, bad = 0, []


def check(t, c, e=""):
    global ok
    if c: ok += 1; print(f"OK   {t} {e}")
    else: bad.append(t); print(f"FAIL {t} {e}")


from docx import Document
from app import create_app
from models import db
from models.user import User
from models.procurement_demand import ProcurementDemand as D
from services.auth import hash_pw
from services import demand_doc

app = create_app(); app.app_context().push()
assert "pms.test.db" in app.config["SQLALCHEMY_DATABASE_URI"], "保险丝：不是测试库"

PKGS = [
    {"预算金额": 500000, "最高限价": 480000, "评审方法": "综合评分法",
     "定价方式": "固定单价", "是否支持联合体投标": "否", "是否允许合同分包": "否",
     "技术要求": "甲包技术ALPHA", "商务要求": "甲包商务ALPHA",
     "特殊资格要求": "甲包资格ALPHA",
     "评审因素": {"价格分": {"分值": 40, "客观项": "是", "标准": "甲包价格标准"},
                "技术要求": {"分值": 50, "客观项": "是", "标准": "甲包技术标准"}}},
    {"预算金额": 356000, "最高限价": 350000, "评审方法": "最低评标价法",
     "定价方式": "固定总价", "是否支持联合体投标": "是", "是否允许合同分包": "是",
     "技术要求": "乙包技术BETA", "商务要求": "乙包商务BETA",
     "特殊资格要求": "乙包资格BETA",
     "评审因素": {"价格分": {"分值": 30, "客观项": "否", "标准": "乙包价格标准"}}},
    {"预算金额": 100000, "评审方法": "综合评分法", "技术要求": "丙包技术GAMMA",
     "商务要求": "丙包商务GAMMA", "特殊资格要求": "丙包资格GAMMA", "评审因素": {}},
]

d = D(demand_type="gov", status="草稿", created_by="验收",
      project_name="分包验收项目", demand_dept="检验科", manage_dept="医学装备部",
      budget_amount=956000, category="服务", year="2026",
      package_split="分包采购", package_count=len(PKGS),
      packages_json=json.dumps(PKGS, ensure_ascii=False),
      items_json=json.dumps([
          {"包号": "1", "编号": "1-1", "标的名称": "甲标的", "数量": "2", "计量单位": "台"},
          {"包号": "2", "编号": "2-1", "标的名称": "乙标的", "数量": "5", "计量单位": "套"},
          {"包号": "3", "编号": "3-1", "标的名称": "丙标的", "数量": "1", "计量单位": "批"},
      ], ensure_ascii=False))
db.session.add(d); db.session.commit()
did = d.id
print(f"造了 3 个包的需求 #{did}\n")

# ═══ 上下文 ═════════════════════════════════════════════════════
ctx = demand_doc.build_context_for(d, None)
pk = ctx.get("分包") or []
check("① 分包上下文有 3 个包", len(pk) == 3, f"{len(pk)} 个")
check("① 包号是中文序号", [x["序号中文"] for x in pk] == ["一", "二", "三"],
      f"{[x['序号中文'] for x in pk]}")
check("① 金额带千分位", pk[0]["预算金额"] == "500,000.00", pk[0]["预算金额"])
check("① 每个包的技术要求各是各的",
      [x["技术要求"] for x in pk] == ["甲包技术ALPHA", "乙包技术BETA", "丙包技术GAMMA"])
check("① 评审因素按包取", pk[0]["评审_价格分_分值"] == 40 and pk[1]["评审_价格分_分值"] == 30,
      f"{pk[0]['评审_价格分_分值']} / {pk[1]['评审_价格分_分值']}")
check("① 没填的评审格是空的", pk[2]["评审_价格分_分值"] == "")

# ═══ 出稿 ═══════════════════════════════════════════════════════
buf, _ = demand_doc.render(d, None)
out = Document(_io.BytesIO(buf.getvalue()))
t = out.tables[0]
rows = ["".join(c.text for c in r.cells) for r in t.rows]
blob = "\n".join(rows)

# 只数「以合同包开头的行」——资格要求条款里也有「合同包」三个字，
# 数出现次数会把它们算进去（第一次跑出来说有 19 段）。
n_pkg = sum(1 for r in rows if r.strip().startswith("合同包"))
check("② 成稿里三个包各出一段", n_pkg == 3, f"{n_pkg} 段")
for cn_ in ("一", "二", "三"):
    check(f"② 出现「合同包{cn_}」", f"合同包{cn_}" in blob)
for k in ("甲包技术ALPHA", "乙包技术BETA", "丙包技术GAMMA",
          "甲包商务ALPHA", "乙包商务BETA", "丙包商务GAMMA",
          "甲包资格ALPHA", "乙包资格BETA", "丙包资格GAMMA"):
    check(f"② 成稿里有 {k}", k in blob)
check("② 评审标准按包出", "甲包价格标准" in blob and "乙包价格标准" in blob)
check("② 没有残留模板语法", "{{" not in blob and "{%" not in blob)

# 包与包之间不串味：包一那段里不该出现包二的内容
i1 = blob.index("合同包一"); i2 = blob.index("合同包二"); i3 = blob.index("合同包三")
seg1, seg2 = blob[i1:i2], blob[i2:i3]
check("② 包一那段里没有包二的内容",
      "乙包技术BETA" not in seg1 and "乙包资格BETA" not in seg1)
check("② 包二那段里没有包一的内容",
      "甲包技术ALPHA" not in seg2 and "甲包资格ALPHA" not in seg2)

# ⑧ 一般资格要求是固定的 8 条，跟着每个包印
check("③ 一般资格要求 8 条都在",
      all(x in blob for x in ("具有独立承担民事责任的能力",
                              "具有良好的商业信誉",
                              "不属于为本项目提供整体设计")),
      "")
check("③ 一般/特殊两节都有",
      "第一节：供应商一般资格要求" in blob and "第二节：供应商特殊资格要求" in blob)

# ═══ 单包与老数据 ═══════════════════════════════════════════════
d2 = D(demand_type="gov", status="草稿", created_by="验收",
       project_name="单包验收项目", demand_dept="检验科", manage_dept="医学装备部",
       budget_amount=88000, category="货物", year="2026",
       package_split="不分包采购",
       tech_requirements="老数据技术要求", business_requirements="老数据商务要求",
       qualification_requirements="老数据资格要求", eval_method="综合评分法")
db.session.add(d2); db.session.commit()
ctx2 = demand_doc.build_context_for(d2, None)
check("④ 没有 packages_json 的老数据也兜出一个包", len(ctx2["分包"]) == 1)
buf2, _ = demand_doc.render(d2, None)
blob2 = "\n".join("".join(c.text for c in r.cells)
                  for r in Document(_io.BytesIO(buf2.getvalue())).tables[0].rows)
check("④ 老数据的技术要求还出得来", "老数据技术要求" in blob2)
_rows2 = [ "".join(c.text for c in r.cells)
           for r in Document(_io.BytesIO(buf2.getvalue())).tables[0].rows ]
_n2 = sum(1 for r in _rows2 if r.strip().startswith("合同包"))
check("④ 老数据只出一个包", _n2 == 1, f"{_n2} 段")

# ═══ 接口：存得下、读得回 ═══════════════════════════════════════
u = db.session.execute(db.select(User).filter_by(username="acc_pkgofficer")).scalar_one_or_none()
if u is None:
    u = User(username="acc_pkgofficer", display_name="验收", role="officer",
             dept_code="", active=1, agency_code="", salt="", pw_hash="")
    db.session.add(u)
u.display_name, u.role, u.active = "验收", "officer", 1
salt = secrets.token_hex(16); u.salt, u.pw_hash = salt, hash_pw(PW, salt)
if hasattr(u, "must_change_pw"): u.must_change_pw = 0
d.created_by = "验收"
db.session.commit()

c = app.test_client()
c.post("/api/auth/login", json={"username": "acc_pkgofficer", "password": PW})
NEW = PKGS + [{"预算金额": 1, "技术要求": "丁包技术DELTA", "评审因素": {}}]
r = c.put(f"/api/procurement-demands/{did}",
          json={"packages_json": json.dumps(NEW, ensure_ascii=False), "package_count": 4})
check("⑤ 分包能存进去", r.status_code == 200, f"HTTP {r.status_code}")
db.session.expire_all()
saved = json.loads(db.session.get(D, did).packages_json)
check("⑤ 存下来是 4 个包", len(saved) == 4, f"{len(saved)} 个")
r = c.get(f"/api/procurement-demands/{did}")
back = (r.get_json() or {}).get("data") or {}
check("⑤ 读得回来", "packages_json" in back and len(json.loads(back["packages_json"])) == 4)

buf3, _ = demand_doc.render(db.session.get(D, did), None)
blob3 = "\n".join("".join(c_.text for c_ in r_.cells)
                  for r_ in Document(_io.BytesIO(buf3.getvalue())).tables[0].rows)
check("⑤ 加的第四个包也出稿了", "丁包技术DELTA" in blob3 and "合同包四" in blob3)

for x in (did, d2.id):
    row = db.session.get(D, x)
    if row: db.session.delete(row)
db.session.commit()
print(f"   已清理测试需求")

print(f"\n通过 {ok} 项" + (f"，失败 {len(bad)}：{bad}" if bad else "，全部通过"))
sys.exit(1 if bad else 0)
