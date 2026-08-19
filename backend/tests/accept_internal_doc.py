"""院内竞选需求也用同一份采购需求表模板出稿的验收。

字段名和政府采购那张对不上，是否题还存成 0/1，所以单独验一遍映射对不对。
"""
import io as _io
import os, sys, secrets
sys.path.insert(0, ".")
os.environ["PMS_CAPTCHA_ON"] = "0"
os.environ["PMS_DB_PATH"] = "/home/huangxb/pms/pms.test.db"
PW = "Ibd!2026"
ok, bad = 0, []


def check(t, c, e=""):
    global ok
    if c: ok += 1; print(f"OK   {t} {e}")
    else: bad.append(t); print(f"FAIL {t} {e}")


from docx import Document
from app import create_app
from models import db
from models.user import User
from models.internal_bid_demand import InternalBidDemand
from services.auth import hash_pw
from services import demand_doc

app = create_app(); app.app_context().push()
assert "pms.test.db" in app.config["SQLALCHEMY_DATABASE_URI"], "保险丝：不是测试库"

d = InternalBidDemand(
    project_name="验收用院内竞选项目", demand_dept="康复科", manage_dept="医学装备部",
    budget_amount=98000, category="货物", year="2026",
    proc_method="院内竞选", pkg_split="分包采购",
    multi_year=1,                 # 是
    actual_settlement=0,          # 否
    perf_bond=1, quality_bond=0,
    invite_expert=1, invite_supplier=0,
    accept_org="委托采购代理机构验收",
    contract_period="2026年1月起；2026年12月止",
    items_json='[{"包号":"1","编号":"1-1","标的名称":"康复训练器","数量":"3","计量单位":"套"}]',
    status="初稿", created_by="验收")
db.session.add(d); db.session.commit()
did = d.id
print(f"造了院内竞选需求 #{did}\n")

ctx = demand_doc.build_context_for(d, None)
check("① 走的是院内竞选那套映射", ctx.get("采购组织形式") == "自行采购", f"{ctx.get('采购组织形式')}")
check("① 0/1 转成了是否", ctx.get("是否属于一签多年项目") == "是"
      and ctx.get("是否为据实结算") == "否",
      f"一签多年={ctx.get('是否属于一签多年项目')} 据实结算={ctx.get('是否为据实结算')}")
check("① 质量保证金 0 也认成否", ctx.get("成交供应商是否需要缴纳质量保证金") == "否")
check("① 履约保证金 1 认成是", ctx.get("成交供应商是否需要缴纳履约保证金") == "是")

buf, missing = demand_doc.render(d, None)
check("② 能出稿", buf.getbuffer().nbytes > 8000, f"{buf.getbuffer().nbytes} 字节")

out = Document(_io.BytesIO(buf.getvalue()))
text = "\n".join(p.text for p in out.paragraphs)
for t in out.tables:
    for row in t.rows:
        for c in row.cells:
            text += "\n" + c.text

check("② 没有残留模板语法", "{{" not in text and "{%" not in text)
check("② 项目与科室填进去了",
      "验收用院内竞选项目" in text and "康复科" in text)
check("② 3.1 勾了自行采购", "■ 自行采购" in text, "")
check("② 3.2.3 勾了院内竞选", "■ 院内竞选" in text)
check("② 3.4 勾了分包采购", "■ 分包采购" in text and "□ 不分包采购" in text)
check("② 验收组织方式勾对", "■ 委托采购代理机构验收" in text)
check("② 标的出来了", "康复训练器" in text and "套" in text)
check("② 合同履行期限填了", "2026年1月起" in text)

# ── 接口 ──────────────────────────────────────────────────────
u = db.session.execute(db.select(User).filter_by(username="acc_ibdofficer")).scalar_one_or_none()
if u is None:
    u = User(username="acc_ibdofficer", display_name="验收", role="officer",
             dept_code="", active=1, agency_code="", salt="", pw_hash="")
    db.session.add(u)
u.display_name, u.role, u.active = "验收", "officer", 1
salt = secrets.token_hex(16); u.salt, u.pw_hash = salt, hash_pw(PW, salt)
if hasattr(u, "must_change_pw"): u.must_change_pw = 0
db.session.commit()

c = app.test_client()
c.post("/api/auth/login", json={"username": "acc_ibdofficer", "password": PW})

r = c.get(f"/api/internal-bid-demands/{did}/doc-status")
dd = (r.get_json() or {}).get("data") or {}
check("③ 出稿状态接口", r.status_code == 200 and dd.get("total", 0) > 40,
      f"已填 {dd.get('filled')}/{dd.get('total')}")

r = c.get(f"/api/internal-bid-demands/{did}/doc", query_string={"download": "1"})
check("③ 下载 Word", r.status_code == 200 and len(r.data) > 8000, f"HTTP {r.status_code}")

r = c.get(f"/api/internal-bid-demands/{did}/doc")
check("③ 预览是 PDF", r.status_code == 200
      and (r.headers.get("Content-Type") or "").startswith("application/pdf"),
      f"HTTP {r.status_code} 类型={r.headers.get('Content-Type')} {len(r.data)} 字节")

db.session.delete(db.session.get(InternalBidDemand, did))
db.session.commit()
print(f"   已清理 #{did}")

print(f"\n通过 {ok} 项" + (f"，失败 {len(bad)}：{bad}" if bad else "，全部通过"))
sys.exit(1 if bad else 0)
