"""采购需求模板库验收（黄新博 2026-08-19 ⑩ 第 1 条）。

「每个经办人可以让其他人 COPY 过去，或者是授权使用」——
所以重点验**授权**：没授权的用不了，授权了才能用；改和删只有主人能做。
"""
import json, os, secrets, sys
sys.path.insert(0, ".")
os.environ["PMS_CAPTCHA_ON"] = "0"
os.environ["PMS_DB_PATH"] = "/home/huangxb/pms/pms.test.db"
PW = "Tpl!2026"
ok, bad = 0, []


def check(t, c, e=""):
    global ok
    if c: ok += 1; print(f"OK   {t} {e}")
    else: bad.append(t); print(f"FAIL {t} {e}")


from app import create_app
from models import db
from models.user import User
from models.procurement_demand import ProcurementDemand as D
from models.demand_template import DemandTemplate as T
from services.auth import hash_pw

app = create_app(); app.app_context().push()
assert "pms.test.db" in app.config["SQLALCHEMY_DATABASE_URI"], "保险丝：不是测试库"


def officer(username, name):
    u = db.session.execute(db.select(User).filter_by(username=username)).scalar_one_or_none()
    if u is None:
        u = User(username=username, display_name=name, role="officer", dept_code="",
                 active=1, agency_code="", salt="", pw_hash="")
        db.session.add(u)
    u.display_name, u.role, u.active = name, "officer", 1
    salt = secrets.token_hex(16); u.salt, u.pw_hash = salt, hash_pw(PW, salt)
    if hasattr(u, "must_change_pw"): u.must_change_pw = 0
    db.session.commit()
    c = app.test_client()
    c.post("/api/auth/login", json={"username": username, "password": PW})
    return c


A = officer("tpl_a", "甲经办人")
B = officer("tpl_b", "乙经办人")

# 甲的需求：填满第三、六~十一部分 + 两个包
src = D(demand_type="gov", status="草稿", created_by="甲经办人",
        project_name="模板源项目", demand_dept="检验科", manage_dept="医学装备部",
        budget_amount=100000, category="货物", year="2026",
        org_form="分散采购", budget_method="公开招标", package_split="分包采购",
        sme_policy="不专门面向中小企业采购",
        business_requirements="源商务要求XYZ", qualification_requirements="源资格要求XYZ",
        eval_method="综合评分法", contract_type="买卖合同",
        payment_terms="源付款约定XYZ", acceptance_org="自行验收",
        packages_json=json.dumps([{"预算金额": 6, "技术要求": "源包一技术"},
                                  {"预算金额": 7, "技术要求": "源包二技术"}],
                                 ensure_ascii=False),
        package_count=2)
db.session.add(src); db.session.commit()
sid = src.id

# 乙的空需求，用来套模板
dst = D(demand_type="gov", status="草稿", created_by="乙经办人",
        project_name="待套用项目", demand_dept="检验科", manage_dept="医学装备部",
        category="货物", year="2026")
db.session.add(dst); db.session.commit()
tid_dst = dst.id
print(f"甲的源需求 #{sid}，乙的空需求 #{tid_dst}\n")

# ═══ 存模板 ═════════════════════════════════════════════════════
r = A.post(f"/api/procurement-demands/{sid}/save-as-template",
           json={"name": "医用耗材类·公开招标", "note": "常规耗材项目用这个"})
d = r.get_json() or {}
check("① 能存成模板", r.status_code == 200, d.get("message", ""))
tpl_id = (d.get("data") or {}).get("id")
check("① 模板归甲所有", (d.get("data") or {}).get("owner") == "甲经办人")
check("① 默认不公开", (d.get("data") or {}).get("shared") is False)

r = A.post(f"/api/procurement-demands/{sid}/save-as-template", json={"name": ""})
check("① 不给名字会拦住", r.status_code == 400, (r.get_json() or {}).get("error", "")[:20])

# ═══ 授权：乙一开始看不到、用不了 ═══════════════════════════════
r = B.get("/api/procurement-demands/templates")
ids = [x["id"] for x in (r.get_json() or {}).get("data") or []]
check("② 没授权时乙看不到这份模板", tpl_id not in ids, f"乙能看到 {len(ids)} 份")

r = B.post(f"/api/procurement-demands/{tid_dst}/apply-template/{tpl_id}")
check("② 没授权时乙套不了", r.status_code == 403,
      (r.get_json() or {}).get("error", "")[:24])

# 甲授权给乙
r = A.put(f"/api/procurement-demands/templates/{tpl_id}",
          json={"shared_with": ["乙经办人"]})
check("② 甲能授权", r.status_code == 200,
      f"{(r.get_json() or {}).get('data', {}).get('shared_with')}")

r = B.get("/api/procurement-demands/templates")
ids = [x["id"] for x in (r.get_json() or {}).get("data") or []]
check("② 授权后乙看得到了", tpl_id in ids)

# ═══ 套用 ═══════════════════════════════════════════════════════
r = B.post(f"/api/procurement-demands/{tid_dst}/apply-template/{tpl_id}")
d = r.get_json() or {}
check("③ 乙能套用了", r.status_code == 200, d.get("message", ""))
db.session.expire_all()
got = db.session.get(D, tid_dst)
check("③ 第六部分套过来了", got.business_requirements == "源商务要求XYZ")
check("③ 第七部分套过来了", got.qualification_requirements == "源资格要求XYZ")
check("③ 第九部分套过来了", got.payment_terms == "源付款约定XYZ")
check("③ 第三部分套过来了", got.budget_method == "公开招标")
check("③ 分包整份套过来了",
      "源包一技术" in (got.packages_json or "") and "源包二技术" in (got.packages_json or ""),
      f"包数={got.package_count}")
check("③ 项目名称没被覆盖（那是本项目独有的）",
      got.project_name == "待套用项目", got.project_name)

r = A.get("/api/procurement-demands/templates")
mine = next(x for x in (r.get_json() or {}).get("data") if x["id"] == tpl_id)
check("③ 用过一次记了次数", mine["use_count"] == 1, f"{mine['use_count']}")

# ═══ 改和删只有主人能做 ═════════════════════════════════════════
r = B.put(f"/api/procurement-demands/templates/{tpl_id}", json={"name": "乙改的名字"})
check("④ 乙改不了甲的模板", r.status_code == 403, (r.get_json() or {}).get("error", "")[:20])
r = B.delete(f"/api/procurement-demands/templates/{tpl_id}")
check("④ 乙删不了甲的模板", r.status_code == 403)

# ═══ 设成公开，谁都能用 ═════════════════════════════════════════
r = A.put(f"/api/procurement-demands/templates/{tpl_id}", json={"shared": True})
check("⑤ 能设成公开", (r.get_json() or {}).get("data", {}).get("shared") is True)
C = officer("tpl_c", "丙经办人")
r = C.get("/api/procurement-demands/templates")
ids = [x["id"] for x in (r.get_json() or {}).get("data") or []]
check("⑤ 公开后没授权过的丙也能用", tpl_id in ids)

# ═══ 上传模板即体检 ═════════════════════════════════════════════
import io as _io
from docx import Document
bad_doc = Document()
bad_doc.add_paragraph("资格审查：☑ 甲方 ☐ 乙方")
bad_doc.add_paragraph("{%tr for r in 标的 %}")
buf = _io.BytesIO(); bad_doc.save(buf); buf.seek(0)
r = A.post("/api/procurement-demands/template-lint",
           data={"file": (buf, "坏模板.docx")}, content_type="multipart/form-data")
d = (r.get_json() or {}).get("data") or {}
check("⑥ 上传坏模板当场报错", r.status_code == 200 and d.get("passed") is False,
      f"{len(d.get('errors') or [])} 个错")
check("⑥ 报的是人话",
      any("字形" in x for x in d.get("errors") or [])
      and any("不成对" in x for x in d.get("errors") or []),
      f"{(d.get('errors') or [''])[0][:40]}")

r = A.post("/api/procurement-demands/template-lint",
           data={"file": (_io.BytesIO(b"xx"), "a.txt")}, content_type="multipart/form-data")
check("⑥ 非 docx 被拒", r.status_code == 400)

for x in (sid, tid_dst):
    row = db.session.get(D, x)
    if row: db.session.delete(row)
row = db.session.get(T, tpl_id)
if row: db.session.delete(row)
for t2 in db.session.execute(db.select(T).filter_by(owner="甲经办人")).scalars():
    db.session.delete(t2)
db.session.commit()
print("   已清理")

print(f"\n通过 {ok} 项" + (f"，失败 {len(bad)}：{bad}" if bad else "，全部通过"))
sys.exit(1 if bad else 0)
